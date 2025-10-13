#!/usr/bin/env python3
"""
enhanced_checkAPI_fast.py
Faster version of your GCI prediction pipeline:
 - Combined GEE yearly fetch (single server-side composite + reduceRegion)
 - JSON caching (.gee_cache)
 - Retry/backoff for GEE calls
 - Speed tuned ML defaults (fewer trees/boosting iters, fewer bootstraps)
 - Keeps original functionality / report generation
"""

import math
import csv
import os
import json
import time
import hashlib
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from geopy.distance import geodesic
from docx import Document
from docx.shared import Inches
import warnings
import requests
warnings.filterwarnings('ignore')

# ML Models
from sklearn.linear_model import LinearRegression
from sklearn.svm import SVR
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import cross_val_score
from enhanced_model_fast import calculate_risk_metrics

# Try faster gradient boosting if available
try:
    from sklearn.ensemble import HistGradientBoostingRegressor
    HAVE_HISTGB = True
except Exception:
    HAVE_HISTGB = False

# Time Series Models
try:
    from statsmodels.tsa.arima.model import ARIMA
    ARIMA_AVAILABLE = True
except Exception:
    ARIMA_AVAILABLE = False

# Deep Learning
try:
    from tensorflow.keras.models import Sequential, Model
    from tensorflow.keras.layers import Dense, LSTM, Dropout, Conv1D, MaxPooling1D, Flatten, Input, Concatenate, BatchNormalization
    from tensorflow.keras.optimizers import Adam, RMSprop
    from tensorflow.keras.callbacks import EarlyStopping
    DL_AVAILABLE = True
except Exception:
    DL_AVAILABLE = False

# Earth Engine (optional)
try:
    import ee
    EE_AVAILABLE = True
except Exception:
    EE_AVAILABLE = False

# -------------------------
# Config / Speed Defaults
# -------------------------
CACHE_DIR = ".gee_cache"
os.makedirs(CACHE_DIR, exist_ok=True)

DEFAULT_RF_TREES = 100
DEFAULT_GB_ITERS = 100
N_BOOTSTRAP = 20
MIN_DL_SAMPLES = 200  # require many samples before running DL
DL_EPOCHS = 20
DL_BATCH = 16

# -------------------------
# EE init + helpers
# -------------------------
def initialize_earth_engine(interactive=True, project='crisiscompass'):
    if not EE_AVAILABLE:
        print("[WARN] Earth Engine Python API not installed.")
        return False
    try:
        ee.Initialize(project=project)
        print("[INFO] Earth Engine initialized.")
        return True
    except Exception as e:
        print(f"[INFO] ee.Initialize() failed: {e}")
        if interactive:
            try:
                ee.Authenticate()
                ee.Initialize(project=project)
                print("[INFO] Earth Engine authenticated and initialized.")
                return True
            except Exception as e2:
                print(f"[ERROR] Earth Engine authentication failed: {e2}")
                return False
        return False

def _try_collections(candidates):
    """Return first working ee.ImageCollection or (None, None)."""
    if not EE_AVAILABLE:
        return None, None
    for cid in candidates:
        try:
            ic = ee.ImageCollection(cid)
            # small metadata call to validate access
            size = ic.size().getInfo()
            # print debug
            print(f"[INFO] Using collection id: {cid} (size={size})")
            return ic, cid
        except Exception as e:
            print(f"[DEBUG] Candidate failed: {cid} -> {e}")
            continue
    print("[ERROR] No candidate collection IDs were available/accessible.")
    return None, None

def _try_images(candidates):
    """Return first working ee.Image or (None, None)."""
    if not EE_AVAILABLE:
        return None, None
    for aid in candidates:
        try:
            img = ee.Image(aid)
            bn = img.bandNames().getInfo()
            print(f"[INFO] Using image id: {aid} (bands: {bn})")
            return img, aid
        except Exception as e:
            print(f"[DEBUG] Candidate image failed: {aid} -> {e}")
            continue
    print("[ERROR] No candidate image IDs were available/accessible.")
    return None, None

def _pick_band_by_keywords(image_or_img, keywords=('temp','t2m','temperature','no2','ndvi','precip','precipitation','built','urban')):
    """Return first band name matching keywords or first band. Accepts ee.Image or ee.ImageCollection."""
    if not EE_AVAILABLE:
        return None
    try:
        # try image_or_img.bandNames()
        if hasattr(image_or_img, 'bandNames'):
            bn = image_or_img.bandNames().getInfo()
        else:
            # assume ImageCollection: take first image
            bn = image_or_img.first().bandNames().getInfo()
    except Exception as e:
        print("[WARN] Could not get band names:", e)
        return None
    band_names_lower = [b.lower() for b in bn]
    for kw in keywords:
        for i, b in enumerate(band_names_lower):
            if kw in b:
                return bn[i]
    return bn[0] if bn else None

# -------------------------
# Caching & retry wrapper
# -------------------------
def _cache_key(lat, lon, year):
    key = f"{lat:.6f}_{lon:.6f}_{year}"
    return os.path.join(CACHE_DIR, hashlib.sha1(key.encode()).hexdigest() + ".json")

def _save_cache(keypath, data):
    try:
        with open(keypath, "w") as f:
            json.dump(data, f)
    except Exception as e:
        print("[WARN] cache save failed:", e)

def _load_cache(keypath):
    try:
        with open(keypath, "r") as f:
            return json.load(f)
    except Exception:
        return None

def retry_with_backoff(func, tries=4, initial_delay=1.0, backoff=2.0):
    def wrapper(*args, **kwargs):
        delay = initial_delay
        for i in range(tries):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                print(f"[WARN] {func.__name__} failed (attempt {i+1}/{tries}): {e}")
                time.sleep(delay)
                delay *= backoff
        raise RuntimeError(f"{func.__name__} failed after {tries} attempts")
    return wrapper

# -------------------------
# Combined yearly fetch (fast)
# -------------------------
@retry_with_backoff
def fetch_yearly_features(lat, lon, year,
                          scale_temp=10000, scale_precip=5000, scale_ndvi=250,
                          scale_no2=1000, scale_ghsl=100):
    """
    Fetch temperature, precipitation, ndvi, no2, urban fraction for a point+year using ONE combined reduceRegion.
    Returns dict with None where data is not available.
    Caches results on disk.
    """
    cache_path = _cache_key(lat, lon, year)
    cached = _load_cache(cache_path)
    if cached:
        # convert JSON numbers back
        return {k: (None if cached.get(k) is None else float(cached.get(k))) for k in cached}

    if not EE_AVAILABLE:
        print("[WARN] Earth Engine not available - cannot fetch real data.")
        res = {'temperature': None, 'precipitation': None, 'ndvi': None, 'no2_concentration': None, 'urban_fraction': None}
        _save_cache(cache_path, res)
        return res

    geom_point = ee.Geometry.Point(lon, lat)
    start = f"{int(year)}-01-01"
    end = f"{int(year)}-12-31"

    # ERA5 / ERA5-Land candidates (temperature)
    era_ids = ["ECMWF_ERA5_LAND_DAILY_AGGR", "ECMWF/ERA5_LAND/DAILY", "ECMWF/ERA5/DAILY"]
    era_ic, _ = _try_collections(era_ids)
    t_img = None
    if era_ic:
        try:
            t_img = era_ic.filterDate(start, end).mean()
            t_band = _pick_band_by_keywords(t_img, ('t2m','temp','temperature','air_temperature'))
            if t_band:
                t_img = t_img.select([t_band]).rename(['temperature'])
        except Exception as e:
            print("[DEBUG] ERA5 reduction error:", e)
            t_img = None

    # CHIRPS precipitation candidates
    chirps_ids = ["UCSB-CHG_CHIRPS_DAILY", "UCSB-CHG/CHIRPS_DAILY", "UCSB-CHG/CHIRPS/DAILY"]
    chirps_ic, _ = _try_collections(chirps_ids)
    p_img = None
    if chirps_ic:
        try:
            p_img = chirps_ic.filterDate(start, end).sum()
            p_band = _pick_band_by_keywords(p_img, ('precip','precipitation','ppt','rain'))
            if p_band:
                p_img = p_img.select([p_band]).rename(['precipitation'])
        except Exception as e:
            print("[DEBUG] CHIRPS reduction error:", e)
            p_img = None

    # MODIS NDVI
    modis_ids = ["MODIS/061/MOD13Q1", "MODIS/006/MOD13Q1", "MODIS/006/MYD13Q1"]
    modis_ic, _ = _try_collections(modis_ids)
    ndvi_img = None
    if modis_ic:
        try:
            ndvi_img = modis_ic.filterDate(start, end).mean()
            ndvi_band = _pick_band_by_keywords(ndvi_img, ('ndvi',))
            if ndvi_band:
                ndvi_img = ndvi_img.select([ndvi_band]).rename(['ndvi'])
        except Exception as e:
            print("[DEBUG] MODIS NDVI reduction error:", e)
            ndvi_img = None

    # Sentinel-5P NO2 candidates
    s5_ids = ["COPERNICUS/S5P/OFFL/L3_NO2", "COPERNICUS/S5P/OFFL/L2__NO2", "COPERNICUS/S5P/OFFL/L3__NO2"]
    s5_ic, _ = _try_collections(s5_ids)
    no2_img = None
    if s5_ic:
        try:
            no2_img = s5_ic.filterDate(start, end).mean()
            no2_band = _pick_band_by_keywords(no2_img, ('no2','tropospheric_column_number_density','column_number_density'))
            if no2_band:
                no2_img = no2_img.select([no2_band]).rename(['no2'])
        except Exception as e:
            print("[DEBUG] Sentinel-5P reduction error:", e)
            no2_img = None

    # GHSL: urban built area (single image)
    ghsl_ids = ["JRC/GHSL/P2023A/GHS_BUILT_S", "JRC/GHSL/P2023A/GHS_BUILT_V"]
    ghsl_img, _ = _try_images(ghsl_ids)
    if ghsl_img:
        try:
            ghsl_img = ghsl_img.rename(['built'])
        except Exception:
            pass

    imgs = []
    if t_img: imgs.append(t_img)
    if p_img: imgs.append(p_img)
    if ndvi_img: imgs.append(ndvi_img)
    if no2_img: imgs.append(no2_img)
    if ghsl_img: imgs.append(ghsl_img)

    if not imgs:
        out = {'temperature': None, 'precipitation': None, 'ndvi': None, 'no2_concentration': None, 'urban_fraction': None}
        _save_cache(cache_path, out)
        return out

    combined = ee.Image.cat(imgs)

    # choose an efficient scale (coarsest of inputs)
    scale = int(max(scale_temp, scale_precip, scale_ndvi, scale_no2, scale_ghsl))
    try:
        # try buffer average first (better representative), with tileScale to help memory
        values = combined.reduceRegion(
            reducer=ee.Reducer.mean(),
            geometry=geom_point.buffer(2500),
            scale=scale,
            bestEffort=True,
            maxPixels=1e13,
            tileScale=4
        ).getInfo()
    except Exception as e:
        # fallback to point reduce (faster, less memory)
        try:
            values = combined.reduceRegion(
                reducer=ee.Reducer.mean(),
                geometry=geom_point,
                scale=scale,
                bestEffort=True,
                maxPixels=1e13,
                tileScale=4
            ).getInfo()
        except Exception as e2:
            print("[ERROR] combined reduceRegion failed:", e2)
            values = None

    out = {'temperature': None, 'precipitation': None, 'ndvi': None, 'no2_concentration': None, 'urban_fraction': None}
    if values:
        for k, v in values.items():
            lk = k.lower()
            try:
                raw = float(v)
            except Exception:
                continue
            # temperature: some datasets in K (ERA5), some in C; assume Kelvin if > 200
            if 'temperature' in lk or 't2m' in lk or 'air_temperature' in lk:
                out['temperature'] = raw - 273.15 if raw > 200 else raw
            elif 'precip' in lk or 'ppt' in lk or 'rain' in lk:
                out['precipitation'] = raw
            elif 'ndvi' in lk:
                # MODIS NDVI often scaled by 10000
                out['ndvi'] = (raw / 10000.0) if abs(raw) > 2.0 else raw
            elif 'no2' in lk:
                # leave units as returned (user can interpret); convert small-magnitude to per-m2 if needed
                out['no2_concentration'] = raw
            elif 'built' in lk or 'ghs' in lk:
                mean_built_m2 = raw
                frac = max(0.0, min(1.0, mean_built_m2 / 10000.0))
                out['urban_fraction'] = frac

    # Save to cache (convert None -> null)
    cache_out = {k: (None if out.get(k) is None else float(out.get(k))) for k in out}
    _save_cache(cache_path, cache_out)
    return out

# -------------------------
# Main environmental wrapper used by the pipeline
# -------------------------
def get_environmental_features(lat, lon, city, year):
    """Return dict with temperature (C), precipitation (mm), urban_fraction (0-1), ndvi (-1..1), no2 (raw)."""
    return fetch_yearly_features(lat, lon, int(year))

# -------------------------
# Other utilities (unchanged behavior, but robust)
# -------------------------
def create_enhanced_features(years, coordinates):
    features = []
    for year in years:
        feature_vector = [
            year,
            year - 2015,
            (year - 2015) ** 2,
            np.sin(2 * np.pi * (year - 2015) / 4),
            np.cos(2 * np.pi * (year - 2015) / 4),
        ]
        features.append(feature_vector)
    return np.array(features)

def calculate_polygon_area(vertices):
    R = 6378137
    vertices = [(math.radians(lat), math.radians(lon)) for lat, lon in vertices]
    area = 0
    for i in range(len(vertices)):
        lat1, lon1 = vertices[i]
        lat2, lon2 = vertices[(i + 1) % len(vertices)]
        area += (lon2 - lon1) * (2 + math.sin(lat1) + math.sin(lat2))
    return abs(area) / 2.0 * (R**2)

def parse_csv_lakes(file_path):
    lakes = []
    with open(file_path, "r", newline='', encoding='utf-8') as csv_file:
        reader = csv.DictReader(csv_file)
        for row in reader:
            lat, lon = row.get('Latitude'), row.get('Longitude')
            if lat and lon:
                try:
                    lake = {k: row[k] for k in row.keys()}
                    lake['latitude'] = float(lat)
                    lake['longitude'] = float(lon)
                    lakes.append(lake)
                except ValueError:
                    continue
    return lakes

def is_within_buffer(plot_vertices, lake_coords, lakes, min_buffer=30, max_buffer=75):
    plot_center = (sum([lat for lat, _ in plot_vertices]) / len(plot_vertices),
                   sum([lon for _, lon in plot_vertices]) / len(plot_vertices))
    closest_lake = None
    closest_distance = float('inf')
    for i, lake_coord in enumerate(lake_coords):
        distance = geodesic(plot_center, lake_coord).meters
        if distance < closest_distance:
            closest_distance = distance
            closest_lake = lakes[i]
        if min_buffer <= distance <= max_buffer:
            return True, distance, closest_lake
    return False, closest_distance, closest_lake

def get_plot_points_from_csv(file_path):
    plot_points, plot_points_gci = [], []
    with open(file_path, "r", newline='', encoding='utf-8') as csv_file:
        reader = csv.DictReader(csv_file)
        for row in reader:
            lat, lon = row.get('Latitude'), row.get('Longitude')
            if lat and lon:
                try:
                    lat, lon = float(lat), float(lon)
                    plot_points.append((lat, lon))
                    plot_points_gci.append((lon, lat))
                except ValueError:
                    continue
    return plot_points, plot_points_gci

# -------------------------
# GCI calculation (original behavior)
# -------------------------
def calculate_mock_gci_values(lat, lon, years):
    """Calculates GCI (NDVI) from MODIS per year (no mock randoms). Falls back to None for missing data."""
    gci_values = []
    for year in years:
        try:
            if not EE_AVAILABLE:
                print(f"[WARN] Earth Engine not available for year {year}, returning None for GCI.")
                gci_values.append(None)
                continue
            # Use fetch_yearly_features for NDVI value
            feats = fetch_yearly_features(lat, lon, int(year))
            ndvi = feats.get('ndvi')
            gci_values.append(ndvi if ndvi is not None else None)
        except Exception as e:
            print(f"[ERROR] GCI calculation error for year {year}: {e}")
            gci_values.append(None)
    return gci_values

# -------------------------
# Model building / predictions (speed tuned)
# -------------------------
def enhanced_models(X, y, environmental_features=None):
    """Train quicker models. Return models_dict and scaler."""
    models_dict = {}
    # build enhanced_X
    if environmental_features:
        enhanced_X = []
        for i, year_features in enumerate(X):
            env_feat = environmental_features[i] if i < len(environmental_features) else {}
            feature_vector = list(year_features) + [
                env_feat.get('temperature') if env_feat.get('temperature') is not None else 0.0,
                env_feat.get('precipitation') if env_feat.get('precipitation') is not None else 0.0,
                env_feat.get('urban_fraction') if env_feat.get('urban_fraction') is not None else 0.0,
                env_feat.get('no2_concentration') if env_feat.get('no2_concentration') is not None else 0.0
            ]
            enhanced_X.append(feature_vector)
        enhanced_X = np.array(enhanced_X)
    else:
        enhanced_X = X

    # handle potential NaNs by filling with zeros
    enhanced_X = np.nan_to_num(enhanced_X, nan=0.0, posinf=0.0, neginf=0.0)

    scaler = StandardScaler()
    enhanced_X_scaled = scaler.fit_transform(enhanced_X)

    # y may contain None; remove those entries
    y_arr = np.array(y, dtype=np.float64)
    valid_idx = ~np.isnan(y_arr)
    if valid_idx.sum() < 2:
        raise ValueError("Not enough valid GCI values to train models.")

    X_train = enhanced_X_scaled[valid_idx]
    y_train = y_arr[valid_idx]

    # Linear Regression
    lr = LinearRegression()
    lr.fit(X_train, y_train)
    models_dict['Linear Regression'] = lr

    # SVR (keep but may be slow for many features)
    svr = SVR(kernel='rbf', C=100, gamma=0.1, epsilon=.01)
    svr.fit(X_train, y_train)
    models_dict['Support Vector Machine'] = svr

    # Random Forest (faster / parallel)
    rf = RandomForestRegressor(n_estimators=DEFAULT_RF_TREES, random_state=42, max_depth=10, n_jobs=-1)
    rf.fit(X_train, y_train)
    models_dict['Random Forest'] = rf

    # Gradient Boosting (use faster hist if available)
    if HAVE_HISTGB:
        from sklearn.ensemble import HistGradientBoostingRegressor
        gb = HistGradientBoostingRegressor(max_iter=DEFAULT_GB_ITERS)
    else:
        gb = GradientBoostingRegressor(n_estimators=DEFAULT_GB_ITERS, learning_rate=0.1, max_depth=6)
    gb.fit(X_train, y_train)
    models_dict['Gradient Boosting'] = gb

    # ARIMA if enough points
    if ARIMA_AVAILABLE and len(y_train) >= 4:
        try:
            arima = ARIMA(y_train, order=(2,1,1))
            arima_fit = arima.fit()
            models_dict['ARIMA'] = arima_fit
        except Exception as e:
            print("[WARN] ARIMA training failed:", e)

    # Deep learning: only if many samples and DL available
    if DL_AVAILABLE and len(y_train) >= MIN_DL_SAMPLES:
        try:
            # reshape for CNN-like input: (samples, timesteps=features, channels=1)
            n_samples, n_feats = X_train.shape
            X_dl = X_train.reshape(n_samples, n_feats, 1)

            # CNN simple
            model = Sequential([
                Conv1D(filters=32, kernel_size=3, activation='relu', input_shape=(n_feats,1)),
                BatchNormalization(),
                MaxPooling1D(pool_size=2),
                Dropout(0.2),
                Flatten(),
                Dense(32, activation='relu'),
                Dropout(0.2),
                Dense(1, activation='linear')
            ])
            model.compile(optimizer=Adam(learning_rate=0.001), loss='mse', metrics=['mae'])
            model.fit(X_dl, y_train, epochs=DL_EPOCHS, batch_size=DL_BATCH, verbose=0,
                      callbacks=[EarlyStopping(patience=4, restore_best_weights=True)], validation_split=0.1)
            models_dict['CNN Deep Learning'] = model
        except Exception as e:
            print("[WARN] DL training failed:", e)

    return models_dict, scaler

def make_predictions(models_dict, scaler, future_years, environmental_features=None):
    predictions = {}
    future_X = create_enhanced_features(future_years.flatten(), [])
    # Build last_env_features
    if environmental_features and len(environmental_features) > 0:
        last_env = environmental_features[-1]
    else:
        last_env = {'temperature': 0.0, 'precipitation': 0.0, 'urban_fraction': 0.0, 'no2_concentration': 0.0}

    enhanced_future = []
    for year_features in future_X:
        temp = last_env.get('temperature', 0.0)
        precip = last_env.get('precipitation', 0.0)
        urban = last_env.get('urban_fraction', 0.0)
        no2 = last_env.get('no2_concentration', 0.0)
        fv = list(year_features) + [
            (0.0 if temp is None else temp),
            (0.0 if precip is None else precip),
            (0.0 if urban is None else urban),
            (0.0 if no2 is None else no2)
        ]
        enhanced_future.append(fv)
    enhanced_future = np.array(enhanced_future)
    enhanced_future = np.nan_to_num(enhanced_future, nan=0.0, posinf=0.0, neginf=0.0)
    enhanced_future_scaled = scaler.transform(enhanced_future)

    for name, model in models_dict.items():
        if name == 'ARIMA':
            try:
                pred = model.forecast(steps=len(future_years))
                predictions[name] = np.array(pred)
            except Exception as e:
                print("[WARN] ARIMA prediction failed:", e)
                predictions[name] = np.zeros(len(future_years))
        elif name in ['CNN Deep Learning', 'Hybrid CNN-LSTM']:
            try:
                X_reshaped = enhanced_future_scaled.reshape(enhanced_future_scaled.shape[0], enhanced_future_scaled.shape[1], 1)
                pred = model.predict(X_reshaped).flatten()
                predictions[name] = pred
            except Exception as e:
                print(f"[WARN] DL prediction failed for {name}: {e}")
                predictions[name] = np.zeros(len(future_years))
        else:
            try:
                predictions[name] = model.predict(enhanced_future_scaled)
            except Exception as e:
                print(f"[WARN] Prediction failed for {name}: {e}")
                predictions[name] = np.zeros(len(future_years))
    return predictions

def get_prediction_intervals(models_dict, scaler, X, y, future_years, confidence=0.95, n_bootstrap=N_BOOTSTRAP):
    """Bootstrap-based prediction intervals (reduced bootstrap count)."""
    intervals = {}
    # build X matrix as used for training (enhanced + scaled)
    # X passed here should be the enhanced raw features (not scaled); we'll scale inside
    X_np = np.nan_to_num(X, nan=0.0, posinf=0.0, neginf=0.0)
    y_np = np.array(y, dtype=np.float64)
    valid_idx = ~np.isnan(y_np)
    X_train = X_np[valid_idx]
    y_train = y_np[valid_idx]
    if len(y_train) < 2:
        return intervals

    # build future matrix and scaled version
    future_X = create_enhanced_features(future_years.flatten(), [])
    # Note: for intervals we assume environmental features constant; this function used only for CI display
    # scale using scaler (if scaler present)
    try:
        future_scaled = scaler.transform(future_X) if scaler is not None else future_X
    except Exception:
        # match input dims (fallback)
        future_scaled = np.nan_to_num(future_X, nan=0.0)

    for name, model in models_dict.items():
        if name == 'ARIMA':
            continue
        bootstrap_preds = []
        for _ in range(n_bootstrap):
            idxs = np.random.choice(len(X_train), len(X_train), replace=True)
            Xb = X_train[idxs]
            yb = y_train[idxs]
            try:
                if name == 'Linear Regression':
                    m = LinearRegression()
                elif name == 'Support Vector Machine':
                    m = SVR(kernel='rbf', C=100, gamma=0.1, epsilon=.01)
                elif name == 'Random Forest':
                    m = RandomForestRegressor(n_estimators=int(DEFAULT_RF_TREES/2), random_state=42, n_jobs=-1)
                elif name == 'Gradient Boosting':
                    if HAVE_HISTGB:
                        from sklearn.ensemble import HistGradientBoostingRegressor
                        m = HistGradientBoostingRegressor(max_iter=int(DEFAULT_GB_ITERS/2))
                    else:
                        m = GradientBoostingRegressor(n_estimators=int(DEFAULT_GB_ITERS/2))
                else:
                    continue
                m.fit(Xb, yb)
                p = m.predict(future_scaled)
                bootstrap_preds.append(p)
            except Exception as e:
                print(f"[DEBUG] bootstrap fit failed for {name}: {e}")
                continue
        if bootstrap_preds:
            bp = np.array(bootstrap_preds)
            lower = np.percentile(bp, (1-confidence)/2*100, axis=0)
            upper = np.percentile(bp, (1+confidence)/2*100, axis=0)
            intervals[name] = {'lower': lower, 'upper': upper}
    return intervals

# -------------------------
# Metrics / reports (unchanged)
# -------------------------
def calculate_enhanced_model_metrics(years, gci_values, models_dict, scaler, environmental_features=None):
    X = create_enhanced_features(years, [])
    # augment with env features if present
    if environmental_features:
        enhanced_X = []
        for i, yf in enumerate(X):
            ef = environmental_features[i] if i < len(environmental_features) else {}
            fv = list(yf) + [
                ef.get('temperature') if ef.get('temperature') is not None else 0.0,
                ef.get('precipitation') if ef.get('precipitation') is not None else 0.0,
                ef.get('urban_fraction') if ef.get('urban_fraction') is not None else 0.0,
                ef.get('no2_concentration') if ef.get('no2_concentration') is not None else 0.0
            ]
            enhanced_X.append(fv)
        enhanced_X = np.array(enhanced_X)
    else:
        enhanced_X = X

    enhanced_X = np.nan_to_num(enhanced_X, nan=0.0)
    try:
        enhanced_X_scaled = scaler.transform(enhanced_X)
    except Exception:
        enhanced_X_scaled = enhanced_X

    y = np.array(gci_values, dtype=np.float64)
    valid_idx = ~np.isnan(y)
    y_valid = y[valid_idx]
    X_valid = enhanced_X_scaled[valid_idx]

    metrics_dict = {}
    for name, model in models_dict.items():
        if name == 'ARIMA':
            continue
        try:
            y_pred = model.predict(X_valid)
            mse = mean_squared_error(y_valid, y_pred)
            rmse = np.sqrt(mse)
            metrics_dict[name] = {
                'R²': r2_score(y_valid, y_pred),
                'MAE': mean_absolute_error(y_valid, y_pred),
                'RMSE': rmse,
                'MAPE': np.mean(np.abs((y_valid - y_pred) / (y_valid + 1e-8))) * 100
            }
        except Exception as e:
            print(f"[WARN] Metrics calc failed for {name}: {e}")
            metrics_dict[name] = {'R²': 0, 'MAE': 0, 'RMSE': 0, 'MAPE': 0}
    return metrics_dict

# Visualization helpers (unchanged)
def create_prediction_bar_chart(predictions, future_years, output_path="predictions_comparison.png"):
    models = list(predictions.keys())
    years = future_years.flatten()
    fig, ax = plt.subplots(figsize=(12, 8))
    x = np.arange(len(years))
    if len(models) == 0:
        return None
    width = 0.8 / max(1, len(models))
    colors = plt.cm.Set3(np.linspace(0, 1, len(models)))
    for i, (model_name, pred_values) in enumerate(predictions.items()):
        if model_name == 'ARIMA':
            continue
        ax.bar(x + i * width, pred_values, width, label=model_name, alpha=0.8, color=colors[i])
    ax.set_xlabel('Future Years')
    ax.set_ylabel('Predicted GCI Values')
    ax.set_title('Model Predictions Comparison')
    ax.set_xticks(x + width * (len(models) - 1) / 2)
    ax.set_xticklabels([str(int(year)) for year in years])
    ax.legend()
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_path, dpi=200, bbox_inches='tight')
    plt.close()
    return output_path

def create_combined_predictions_plot(years, gci_values, future_years, predictions, uncertainties, output_path="combined_predictions.png"):
    try:
        plt.figure(figsize=(14, 8))
        # historical
        hist_valid = [(y,g) for y,g in zip(years, gci_values) if g is not None]
        if len(hist_valid) > 0:
            ys = [v[0] for v in hist_valid]
            gs = [v[1] for v in hist_valid]
            plt.plot(ys, gs, 'o-', color='black', label='Historical GCI', linewidth=2)
        # future
        fy = future_years.flatten()
        colors = plt.cm.tab10(np.linspace(0, 1, max(1, len(predictions))))
        markers = ['o','s','^','D','v','>','<','P','X','*']
        for i, (name, pred) in enumerate(predictions.items()):
            if name == 'ARIMA':
                continue
            plt.plot(fy, pred, '-', label=name, color=colors[i % len(colors)], linewidth=2, alpha=0.9, marker=markers[i % len(markers)], markersize=5)
            if name in uncertainties:
                lower = uncertainties[name]['lower']
                upper = uncertainties[name]['upper']
                plt.fill_between(fy, lower, upper, color=colors[i % len(colors)], alpha=0.15)
        plt.title('Combined Model Predictions with Confidence Intervals')
        plt.xlabel('Year'); plt.ylabel('Predicted GCI (NDVI)')
        plt.grid(True, alpha=0.3); plt.legend()
        plt.tight_layout(); plt.savefig(output_path, dpi=200, bbox_inches='tight'); plt.close()
        return output_path
    except Exception as e:
        print('[WARN] Failed combined predictions plot:', e)
        return None

def create_model_performance_bar_chart(metrics_dict, output_path="model_performance.png"):
    models = list(metrics_dict.keys())
    r2_scores = [metrics_dict[model].get('R²', 0) for model in models]
    mae_scores = [metrics_dict[model].get('MAE', 0) for model in models]
    rmse_scores = [metrics_dict[model].get('RMSE', 0) for model in models]

    fig, (ax1, ax2, ax3) = plt.subplots(1, 3, figsize=(15, 5))
    bars1 = ax1.bar(models, r2_scores)
    ax1.set_title('R² Score Comparison'); ax1.set_ylim(0,1)
    for bar, score in zip(bars1, r2_scores):
        ax1.text(bar.get_x()+bar.get_width()/2., score+0.01, f'{score:.3f}', ha='center')
    bars2 = ax2.bar(models, mae_scores)
    ax2.set_title('MAE Comparison')
    for bar, score in zip(bars2, mae_scores):
        ax2.text(bar.get_x()+bar.get_width()/2., score+0.001, f'{score:.3f}', ha='center')
    bars3 = ax3.bar(models, rmse_scores)
    ax3.set_title('RMSE Comparison')
    for bar, score in zip(bars3, rmse_scores):
        ax3.text(bar.get_x()+bar.get_width()/2., score+0.001, f'{score:.3f}', ha='center')
    plt.tight_layout(); plt.savefig(output_path, dpi=200, bbox_inches='tight'); plt.close()
    return output_path

# Report generation uses the existing routines - left largely intact but will cope with None values.

def save_enhanced_report_with_risks(file_path, years, gci_values, future_years, models_dict, 
                                  metrics_dict, plot_points, area, within_buffer, distance, 
                                  closest_lake, risks, environmental_data, uncertainties, scaler,
                                  validation_results=None, output_file="report.docx"):
    doc = Document()
    doc.add_heading("Enhanced Green Cover Index (GCI) Comprehensive Report", level=1)
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph("This enhanced report provides comprehensive analysis of GCI trends, "
                     "model predictions, environmental risk assessment, and uncertainty quantification.")
    # Plot table
    doc.add_heading("Plot Analysis", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Latitude"; hdr[1].text = "Longitude"
    for lat, lon in plot_points:
        row = table.add_row().cells
        row[0].text = f"{lat:.6f}"
        row[1].text = f"{lon:.6f}"
    doc.add_paragraph(f"Calculated area: {area:.2f} m² ({area/10000:.2f} hectares)")
    if within_buffer:
        doc.add_paragraph(f"Plot is within legal buffer, {distance:.2f} m from nearest lake.")
    else:
        doc.add_paragraph("Plot outside legal buffer. Closest lake details:")
        if closest_lake:
            table = doc.add_table(rows=1, cols=2); table.style="Table Grid"
            hdr = table.rows[0].cells; hdr[0].text="Attribute"; hdr[1].text="Value"
            for k, v in closest_lake.items():
                row = table.add_row().cells
                row[0].text = str(k); row[1].text = str(v)
    # Environmental Risk Assessment
    doc.add_heading("Environmental Risk Assessment", level=2)
    risk_table = doc.add_table(rows=1, cols=2); risk_table.style="Table Grid"
    hdr = risk_table.rows[0].cells; hdr[0].text="Risk Category"; hdr[1].text="Risk Level"
    for risk_type, risk_level in risks.items():
        row = risk_table.add_row().cells
        row[0].text = risk_type.replace('_',' ').title(); row[1].text = risk_level
    # Historical GCI table
    doc.add_heading("Historical GCI Data", level=2)
    gci_table = doc.add_table(rows=1, cols=2); gci_table.style="Table Grid"
    hdr = gci_table.rows[0].cells; hdr[0].text="Year"; hdr[1].text="GCI"
    for y, gci in zip(years, gci_values):
        row = gci_table.add_row().cells
        row[0].text = str(y); row[1].text = "N/A" if gci is None else f"{gci:.4f}"
    # Small historical plot (skip if all None)
    try:
        valid = [ (y,g) for y,g in zip(years, gci_values) if g is not None ]
        if len(valid) >= 2:
            ys = [v[0] for v in valid]; gs = [v[1] for v in valid]
            plt.figure(figsize=(8,5)); plt.plot(ys, gs, 'o-'); plt.title("Historical GCI"); plt.xlabel("Year"); plt.ylabel("GCI")
            plt.grid(True); fname="hist_gci.png"; plt.savefig(fname, dpi=200); plt.close()
            doc.add_picture(fname, width=Inches(6)); os.remove(fname)
    except Exception as e:
        print("[WARN] Could not create historical GCI plot:", e)
    # Model metrics summary
    doc.add_heading("Model Performance Metrics", level=2)
    if metrics_dict:
        perf_table = doc.add_table(rows=1, cols=5); perf_table.style="Table Grid"
        hdr = perf_table.rows[0].cells
        hdr[0].text="Model"; hdr[1].text="R²"; hdr[2].text="MAE"; hdr[3].text="RMSE"; hdr[4].text="MAPE (%)"
        for name, m in metrics_dict.items():
            row = perf_table.add_row().cells
            row[0].text = name
            row[1].text = f"{m['R²']:.3f}"
            row[2].text = f"{m['MAE']:.3f}"
            row[3].text = f"{m['RMSE']:.3f}"
            row[4].text = f"{m['MAPE']:.1f}"
    # Predictions with CI tables
    doc.add_heading("Predictions (future)", level=2)
    try:
        preds = make_predictions(models_dict, scaler, future_years, environmental_data)
        for model_name, values in preds.items():
            doc.add_heading(model_name + " Predictions", level=3)
            t = doc.add_table(rows=1, cols=3); t.style="Table Grid"
            hdr = t.rows[0].cells; hdr[0].text="Year"; hdr[1].text="Predicted GCI"; hdr[2].text="CI"
            for i, year in enumerate(future_years.flatten()):
                row = t.add_row().cells
                row[0].text = str(int(year))
                row[1].text = f"{values[i]:.4f}"
                if model_name in uncertainties:
                    lo = uncertainties[model_name]['lower'][i]; up = uncertainties[model_name]['upper'][i]
                    row[2].text = f"[{lo:.4f}, {up:.4f}]"
                else:
                    row[2].text = "N/A"
    except Exception as e:
        print("[WARN] Could not create predictions section:", e)

    # Combined predictions figure with confidence intervals
    try:
        combined_path = create_combined_predictions_plot(years, gci_values, future_years, preds, uncertainties)
        if combined_path:
            doc.add_picture(combined_path, width=Inches(7))
    except Exception as e:
        print('[WARN] Could not add combined predictions figure:', e)

    # Metrics comparison figure and brief comparative text + metrics table near graph
    try:
        perf_path = create_model_performance_bar_chart(metrics_dict)
        if perf_path:
            doc.add_paragraph("\nModel Comparison:")
            doc.add_picture(perf_path, width=Inches(7))
            # quick textual comparative summary
            best_by_r2 = sorted(metrics_dict.items(), key=lambda x: x[1].get('R²', 0), reverse=True)
            if best_by_r2:
                name, m = best_by_r2[0]
                doc.add_paragraph(f"Best R²: {name} (R²={m.get('R²',0):.3f}, MAE={m.get('MAE',0):.3f}, RMSE={m.get('RMSE',0):.3f})")
            # add a concise metrics table adjacent to the graph section
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Model"; hdr[1].text = "R²"; hdr[2].text = "MAE"; hdr[3].text = "RMSE"; hdr[4].text = "MAPE (%)"
            for model_name, mm in metrics_dict.items():
                row = table.add_row().cells
                row[0].text = model_name
                row[1].text = f"{mm.get('R²',0):.3f}"
                row[2].text = f"{mm.get('MAE',0):.3f}"
                row[3].text = f"{mm.get('RMSE',0):.3f}"
                row[4].text = f"{mm.get('MAPE',0):.1f}"
    except Exception as e:
        print('[WARN] Could not add metrics comparison figure:', e)
    doc.save(output_file)
    print(f"[INFO] Enhanced report saved as {output_file}")

# -------------------------
# Main enhanced pipeline
# -------------------------
def enhanced_main(file_path, lake_csv):
    print("Initializing Earth Engine...")
    ee_ok = initialize_earth_engine()
    if not ee_ok:
        print("[WARN] Earth Engine unavailable. This script depends on GEE for real environmental data.")
    print("Processing plot data...")
    plot_points, plot_points_gci = get_plot_points_from_csv(file_path)
    if not plot_points:
        return print("[ERROR] No valid plot points.")
    area = calculate_polygon_area(plot_points)
    print(f"Plot area: {area:.2f} m² ({area/10000:.2f} hectares)")
    print("Processing lake data...")
    lakes = parse_csv_lakes(lake_csv)
    lake_coords = [(lake['latitude'], lake['longitude']) for lake in lakes] if lakes else []
    within_buffer, distance, closest_lake = is_within_buffer(plot_points, lake_coords, lakes) if lakes else (False, float('inf'), None)
    print(f"Buffer analysis: Within buffer = {within_buffer}, Distance = {distance:.2f}m")
    # centroid
    lat = float(np.mean([pt[0] for pt in plot_points])); lon = float(np.mean([pt[1] for pt in plot_points]))
    years = np.arange(2017, 2025)
    print("Fetching GCI (NDVI) for each historical year (cached)...")
    gci_values = calculate_mock_gci_values(lat, lon, years)
    print("Fetching environmental features for each year (cached)...")
    global environmental_data
    environmental_data = []
    for y in years:
        feats = fetch_yearly_features(lat, lon, int(y))
        environmental_data.append({
            'temperature': feats.get('temperature'),
            'precipitation': feats.get('precipitation'),
            'urban_fraction': feats.get('urban_fraction'),
            'ndvi': feats.get('ndvi'),
            'no2_concentration': feats.get('no2_concentration')
        })
    print("GCI values (sample):", gci_values)
    print("Environmental sample (latest):", environmental_data[-1] if environmental_data else None)
    print("Performing risk assessment...")
    risks = calculate_risk_metrics(plot_points_gci, gci_values, environmental_data)
    print("Training models (speed tuned)...")
    X = create_enhanced_features(years, [])
    try:
        models_dict, scaler = enhanced_models(X, gci_values, environmental_data)
    except Exception as e:
        print("[ERROR] Model training failed:", e)
        return
    print("Calculating metrics...")
    metrics_dict = calculate_enhanced_model_metrics(years, gci_values, models_dict, scaler, environmental_data)
    print("Making future predictions...")
    future_years = np.arange(2025, 2036).reshape(-1,1)
    predictions = make_predictions(models_dict, scaler, future_years, environmental_data)
    print("Calculating uncertainties (bootstrap)...")
    uncertainties = get_prediction_intervals(models_dict, scaler, X, gci_values, future_years)
    print("Generating report (may still take time for plotting/reporting)...")
    try:
        save_enhanced_report_with_risks(file_path, years, gci_values, future_years, models_dict,
                                       metrics_dict, plot_points, area, within_buffer, distance,
                                       closest_lake, risks, environmental_data, uncertainties, scaler,
                                       validation_results=None, output_file="report.docx")
    except Exception as e:
        print("[ERROR] Report generation failed:", e)
    print("[INFO] Done.")

# -------------------------
# Run script
# -------------------------
if __name__ == "__main__":
    import sys
    plot_csv = sys.argv[1] if len(sys.argv) > 1 else "coordinates.csv"
    lake_csv = sys.argv[2] if len(sys.argv) > 2 else "output_lake.csv"
    enhanced_main(plot_csv, lake_csv)
