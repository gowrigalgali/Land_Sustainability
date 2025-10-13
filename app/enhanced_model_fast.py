# checkAPI_gee_with_report_images_report_improved.py
import math
import csv
import os
import sys
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from geopy.distance import geodesic
from docx import Document
from docx.shared import Inches
import warnings
import requests
warnings.filterwarnings('ignore')

# ML Models
from sklearn.linear_model import LinearRegression
from sklearn.svm import SVR
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor, VotingRegressor
from sklearn.ensemble import VotingRegressor
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error
from sklearn.preprocessing import MinMaxScaler, StandardScaler
from sklearn.model_selection import cross_val_score, TimeSeriesSplit

# Time Series Models
try:
    from statsmodels.tsa.arima.model import ARIMA
    ARIMA_AVAILABLE = True
except ImportError:
    ARIMA_AVAILABLE = False

# Deep Learning
try:
    from tensorflow.keras.models import Sequential, Model
    from tensorflow.keras.layers import Dense, LSTM, Dropout, Conv1D, MaxPooling1D, Flatten, Input, Concatenate, BatchNormalization
    from tensorflow.keras.optimizers import Adam, RMSprop
    from tensorflow.keras.callbacks import EarlyStopping, ReduceLROnPlateau
    DL_AVAILABLE = True
except ImportError:
    DL_AVAILABLE = False

# Earth Engine (optional)
try:
    import ee
    EE_AVAILABLE = True
except ImportError:
    EE_AVAILABLE = False

# -----------------------------
# Earth Engine initialization (with interactive auth fallback)
# -----------------------------
def initialize_earth_engine(interactive=True):
    if not EE_AVAILABLE:
        print("[WARN] Earth Engine Python package not installed (ee unavailable).")
        return False
    try:
        ee.Initialize(project='crisiscompass')
        print("[INFO] Earth Engine initialized.")
        return True
    except Exception as e:
        print(f"[INFO] ee.Initialize() failed: {e}")
        if interactive:
            try:
                ee.Authenticate()
                ee.Initialize(project='crisiscompass')
                print("[INFO] Earth Engine authenticated and initialized.")
                return True
            except Exception as e2:
                print(f"[ERROR] Earth Engine authentication failed: {e2}")
                return False
        return False

# -----------------------------
# GEE helper utilities (robust loaders / band pickers)
# -----------------------------
def _try_collections(candidates):
    for cid in candidates:
        try:
            ic = ee.ImageCollection(cid)
            size = ic.size().getInfo()
            print(f"[INFO] Using collection id: {cid} (size={size})")
            return ic, cid
        except Exception as e:
            print(f"[DEBUG] Candidate failed: {cid} -> {e}")
            continue
    print("[ERROR] No candidate collection IDs were available/accessible.")
    return None, None

def _try_images(candidates):
    for aid in candidates:
        try:
            img = ee.Image(aid)
            bn = img.bandNames().getInfo()
            print(f"[INFO] Using image id: {aid} (bands: {bn})")
            return img, aid
        except Exception as e:
            print(f"[DEBUG] Candidate image failed: {aid} -> {e}")
            continue
    print("[ERROR] No candidate image IDs were available/accessible.")
    return None, None

def _pick_band_by_keywords_from_image(img, keywords):
    try:
        band_names = img.bandNames().getInfo()
    except Exception as e:
        print("[WARN] cannot read band names:", e)
        return None
    band_names_lower = [b.lower() for b in band_names]
    for kw in keywords:
        for i, bn in enumerate(band_names_lower):
            if kw in bn:
                return band_names[i]
    return band_names[0] if band_names else None

# combined fetch helper (used previously in your file)
def fetch_environmental_features_combined(lat, lon, year, scale=1000):
    if not EE_AVAILABLE:
        return {
            'temperature': None,
            'precipitation': None,
            'ndvi': None,
            'no2_concentration': None,
            'urban_fraction': None
        }

    pt = ee.Geometry.Point(lon, lat)
    start = f"{int(year)}-01-01"
    end = f"{int(year)}-12-31"

    band_images = {}
    try:
        # ERA5 temperature
        temp_cands = ["ECMWF_ERA5_LAND_DAILY_AGGR", "ECMWF/ERA5_LAND/DAILY", "ECMWF/ERA5/DAILY"]
        ic_temp, _ = _try_collections(temp_cands)
        if ic_temp:
            img_temp = ic_temp.filterDate(start, end).mean()
            b = _pick_band_by_keywords_from_image(img_temp, ('t2m','temp','temperature','air_temperature'))
            if b:
                band_images['temperature'] = img_temp.select([b]).rename('temperature')
    except Exception as e:
        print("[DEBUG] Temperature composite prep failed:", e)

    try:
        # CHIRPS precipitation
        prec_cands = ["UCSB-CHG_CHIRPS_DAILY", "UCSB-CHG/CHIRPS_DAILY", "UCSB-CHG/CHIRPS/DAILY"]
        ic_prec, _ = _try_collections(prec_cands)
        if ic_prec:
            img_prec = ic_prec.filterDate(start, end).sum()
            b = _pick_band_by_keywords_from_image(img_prec, ('precip','precipitation','ppt','rain'))
            if b:
                band_images['precipitation'] = img_prec.select([b]).rename('precipitation')
    except Exception as e:
        print("[DEBUG] Precipitation composite prep failed:", e)

    try:
        # MODIS NDVI
        ndvi_cands = ["MODIS/061/MOD13Q1", "MODIS/006/MOD13Q1", "MODIS/006/MYD13Q1"]
        ic_ndvi, _ = _try_collections(ndvi_cands)
        if ic_ndvi:
            img_ndvi = ic_ndvi.filterDate(start, end).mean()
            b = _pick_band_by_keywords_from_image(img_ndvi, ('ndvi',))
            if b:
                band_images['ndvi'] = img_ndvi.select([b]).rename('ndvi')
    except Exception as e:
        print("[DEBUG] NDVI composite prep failed:", e)

    try:
        # Sentinel-5P NO2
        no2_cands = ["COPERNICUS/S5P/OFFL/L3_NO2", "COPERNICUS/S5P/OFFL/L2__NO2"]
        ic_no2, _ = _try_collections(no2_cands)
        if ic_no2:
            img_no2 = ic_no2.filterDate(start, end).mean()
            b = _pick_band_by_keywords_from_image(img_no2, ('no2','column_number_density','tropospheric_column_number_density'))
            if b:
                band_images['no2'] = img_no2.select([b]).rename('no2')
    except Exception as e:
        print("[DEBUG] NO2 composite prep failed:", e)

    try:
        ghsl_img, _ = _try_images(["JRC/GHSL/P2023A/GHS_BUILT_S", "JRC/GHSL/P2023A/GHS_BUILT_V"])
        if ghsl_img:
            band_images['built'] = ghsl_img.rename('built')
    except Exception as e:
        print("[DEBUG] GHSL composite prep failed:", e)

    if not band_images:
        print("[WARN] No GEE datasets available for combined fetch.")
        return {'temperature': None, 'precipitation': None, 'ndvi': None, 'no2_concentration': None, 'urban_fraction': None}

    try:
        composite = ee.Image.cat(list(band_images.values()))
        rr = composite.reduceRegion(reducer=ee.Reducer.mean(), geometry=pt, scale=int(scale), bestEffort=True, maxPixels=1e13).getInfo()
    except Exception as e:
        print("[ERROR] Combined reduceRegion failed:", e)
        return {'temperature': None, 'precipitation': None, 'ndvi': None, 'no2_concentration': None, 'urban_fraction': None}

    out = {'temperature': None, 'precipitation': None, 'ndvi': None, 'no2_concentration': None, 'urban_fraction': None}
    if not rr:
        return out

    if 'temperature' in rr:
        try:
            temp_raw = float(rr['temperature'])
            out['temperature'] = temp_raw - 273.15 if temp_raw > 100 else temp_raw
        except Exception:
            out['temperature'] = None
    if 'precipitation' in rr:
        try:
            out['precipitation'] = float(rr['precipitation'])
        except Exception:
            out['precipitation'] = None
    if 'ndvi' in rr:
        try:
            nd = float(rr['ndvi'])
            if abs(nd) > 2.0:
                nd = nd / 10000.0
            out['ndvi'] = float(nd)
        except Exception:
            out['ndvi'] = None
    if 'no2' in rr:
        try:
            raw_no2 = float(rr['no2'])
            out['no2_concentration'] = raw_no2 * 1e6
        except Exception:
            out['no2_concentration'] = None
    if 'built' in rr:
        try:
            mean_built_m2 = float(rr['built'])
            fract = mean_built_m2 / 10000.0
            out['urban_fraction'] = max(0.0, min(1.0, fract))
        except Exception:
            out['urban_fraction'] = None
    return out

# -----------------------------
# Replace environment aggregator (no mocks)
# -----------------------------
def get_environmental_features(lat, lon, city, year):
    features = fetch_environmental_features_combined(lat, lon, year, scale=1000)
    return features

# -----------------------------
# Feature engineering and helpers
# -----------------------------
def create_enhanced_features(years, coordinates):
    features = []
    for year in years:
        feature_vector = [
            year,
            year - 2015,
            (year - 2015) ** 2,
            np.sin(2 * np.pi * (year - 2015) / 4),
            np.cos(2 * np.pi * (year - 2015) / 4),
        ]
        features.append(feature_vector)
    return np.array(features)

def add_lag_features(gci_values, lags=[1,2]):
    n = len(gci_values)
    out = np.full((n, len(lags)), np.nan)
    for i in range(n):
        for j, lag in enumerate(lags):
            idx = i - lag
            if idx >= 0:
                out[i, j] = gci_values[idx]
    return out

# -----------------------------
# Polygon Area / CSV helpers
# -----------------------------
def calculate_polygon_area(vertices):
    R = 6378137
    vertices = [(math.radians(lat), math.radians(lon)) for lat, lon in vertices]
    area = 0
    for i in range(len(vertices)):
        lat1, lon1 = vertices[i]
        lat2, lon2 = vertices[(i + 1) % len(vertices)]
        area += (lon2 - lon1) * (2 + math.sin(lat1) + math.sin(lat2))
    return abs(area) / 2.0 * (R**2)

def parse_csv_lakes(file_path):
    lakes = []
    with open(file_path, "r", newline='', encoding='utf-8') as csv_file:
        reader = csv.DictReader(csv_file)
        for row in reader:
            lat, lon = row.get('Latitude'), row.get('Longitude')
            if lat and lon:
                try:
                    lake = {k: row[k] for k in row.keys()}
                    lake['latitude'] = float(lat)
                    lake['longitude'] = float(lon)
                    lakes.append(lake)
                except ValueError:
                    continue
    return lakes

def get_plot_points_from_csv(file_path):
    plot_points, plot_points_gci = [], []
    with open(file_path, "r", newline='', encoding='utf-8') as csv_file:
        reader = csv.DictReader(csv_file)
        for row in reader:
            lat, lon = row.get('Latitude'), row.get('Longitude')
            if lat and lon:
                try:
                    lat, lon = float(lat), float(lon)
                    plot_points.append((lat, lon))
                    plot_points_gci.append((lon, lat))
                except ValueError:
                    continue
    return plot_points, plot_points_gci

def is_within_buffer(plot_vertices, lake_coords, lakes, min_buffer=30, max_buffer=75):
    plot_center = (sum([lat for lat, _ in plot_vertices])/len(plot_vertices),
                   sum([lon for _, lon in plot_vertices])/len(plot_vertices))
    closest_lake = None
    closest_distance = float('inf')
    for i, lake_coord in enumerate(lake_coords):
        distance = geodesic(plot_center, lake_coord).meters
        if distance < closest_distance:
            closest_distance = distance
            closest_lake = lakes[i]
        if min_buffer <= distance <= max_buffer:
            return True, distance, closest_lake
    return False, closest_distance, closest_lake

# -----------------------------
# Mock-less GCI calculation (uses GEE)
# -----------------------------
def get_ndvi(lat, lon, year, scale=250):
    candidates = ["MODIS/061/MOD13Q1", "MODIS/006/MOD13Q1", "MODIS/006/MYD13Q1"]
    ic, used = _try_collections(candidates)
    if ic is None:
        print("[WARN] No MODIS NDVI collection found.")
        return None
    col = ic.filterDate(f"{int(year)}-01-01", f"{int(year)}-12-31")
    try:
        year_img = col.mean()
    except Exception:
        try:
            year_img = ee.Image(col.first())
        except Exception as e:
            print("[ERROR] MODIS reduce to image failed:", e)
            return None
    band = _pick_band_by_keywords_from_image(year_img, ('ndvi',))
    if not band:
        print("[WARN] No NDVI band found")
        return None
    try:
        val_dict = year_img.select([band]).reduceRegion(
            reducer=ee.Reducer.mean(),
            geometry=ee.Geometry.Point(lon, lat),
            scale=scale,
            bestEffort=True,
            maxPixels=1e13
        ).getInfo()
        if val_dict and band in val_dict:
            raw = float(val_dict[band])
            if abs(raw) > 2.0:
                return raw / 10000.0
            return raw
    except Exception as e:
        print("[ERROR] reduceRegion failed for NDVI:", e)
    return None

def calculate_gci_values_from_gee(lat, lon, years):
    gci_values = []
    for year in years:
        ndvi_val = get_ndvi(lat, lon, int(year))
        if ndvi_val is None:
            print(f"[WARN] NDVI missing for year {year}")
            gci_values.append(None)
        else:
            gci_values.append(float(ndvi_val))
    return gci_values

# -----------------------------
# Risk assessment (unchanged)
# -----------------------------
def calculate_risk_metrics(coordinates, predictions, environmental_data):
    risks = {}
    valid_predictions = [p for p in predictions if p is not None]
    if len(valid_predictions) < 2:
        risks = {
            'vegetation_decline': 'Medium',
            'climate_vulnerability': 'Medium',
            'urbanization_pressure': 'Medium',
            'air_quality': 'Medium'
        }
        return risks

    trend_slope = np.polyfit(range(len(valid_predictions)), valid_predictions, 1)[0]
    if trend_slope < -0.05:
        risks['vegetation_decline'] = 'High'
    elif trend_slope < 0:
        risks['vegetation_decline'] = 'Medium'
    else:
        risks['vegetation_decline'] = 'Low'

    if environmental_data:
        temp_values = [data.get('temperature') for data in environmental_data if data.get('temperature') is not None]
        if len(temp_values) > 1:
            temp_trend = np.polyfit(range(len(temp_values)), temp_values, 1)[0]
            if temp_trend > 0.5:
                risks['climate_vulnerability'] = 'High'
            elif temp_trend > 0.2:
                risks['climate_vulnerability'] = 'Medium'
            else:
                risks['climate_vulnerability'] = 'Low'
        else:
            risks['climate_vulnerability'] = 'Unknown'
    else:
        risks['climate_vulnerability'] = 'Unknown'

    if environmental_data:
        precip_values = [data.get('precipitation') for data in environmental_data if data.get('precipitation') is not None]
        if len(precip_values) > 1:
            precip_trend = np.polyfit(range(len(precip_values)), precip_values, 1)[0]
            if precip_trend < -0.1:
                risks['water_stress'] = 'High'
            elif precip_trend < 0:
                risks['water_stress'] = 'Medium'
            else:
                risks['water_stress'] = 'Low'
        else:
            risks['water_stress'] = 'Unknown'
    else:
        risks['water_stress'] = 'Unknown'

    if environmental_data:
        urban_values = [data.get('urban_fraction') for data in environmental_data if data.get('urban_fraction') is not None]
        if len(urban_values) > 1:
            urban_growth = np.polyfit(range(len(urban_values)), urban_values, 1)[0]
            if urban_growth > 0.1:
                risks['urbanization_pressure'] = 'High'
            elif urban_growth > 0.05:
                risks['urbanization_pressure'] = 'Medium'
            else:
                risks['urbanization_pressure'] = 'Low'
        else:
            risks['urbanization_pressure'] = 'Unknown'
    else:
        risks['urbanization_pressure'] = 'Unknown'

    if environmental_data:
        no2_values = [data.get('no2_concentration') for data in environmental_data if data.get('no2_concentration') is not None]
        if len(no2_values) > 0:
            avg_no2 = np.mean(no2_values)
            if avg_no2 > 150:
                risks['air_quality'] = 'High'
            elif avg_no2 > 50:
                risks['air_quality'] = 'Medium'
            else:
                risks['air_quality'] = 'Low'
        else:
            risks['air_quality'] = 'Unknown'
    else:
        risks['air_quality'] = 'Unknown'

    return risks

# -----------------------------
# Validation / Models / Predictions (unchanged except n_jobs and handling)
# -----------------------------
def validate_model_performance(model, X, y, model_name):
    validation_results = {}
    try:
        n_samples = len(X)
        n_splits = max(2, min(5, n_samples // 2))
        tscv = TimeSeriesSplit(n_splits=n_splits)
        cv_scores = cross_val_score(model, X, y, cv=tscv, scoring='r2')
        valid_scores = cv_scores[~np.isnan(cv_scores)]
        validation_results['cv_mean'] = float(valid_scores.mean()) if len(valid_scores) > 0 else 0.0
        validation_results['cv_std'] = float(valid_scores.std()) if len(valid_scores) > 0 else 0.0
        train_score = model.score(X, y)
        validation_results['train_score'] = float(train_score)
        validation_results['overfitting_warning'] = False
        if len(valid_scores) > 0 and train_score - valid_scores.mean() > 0.1:
            validation_results['overfitting_warning'] = True
        if hasattr(model, 'n_estimators'):
            validation_results['complexity'] = f"n_estimators: {getattr(model, 'n_estimators', 'NA')}"
        elif hasattr(model, 'C'):
            validation_results['complexity'] = f"C: {getattr(model, 'C', 'NA')}"
        else:
            validation_results['complexity'] = "Standard parameters"

        if len(valid_scores) > 0 and valid_scores.mean() > 0.7 and not validation_results['overfitting_warning']:
            validation_results['status'] = 'Good'
        elif len(valid_scores) > 0 and valid_scores.mean() > 0.3:
            validation_results['status'] = 'Needs Tuning'
        else:
            validation_results['status'] = 'Limited Data'
    except Exception as e:
        validation_results['error'] = str(e)
        validation_results['status'] = 'Error'
        validation_results['cv_mean'] = 0.0
        validation_results['cv_std'] = 0.0
        validation_results['train_score'] = 0.0
    return validation_results

# -----------------------------
# Deep models (unchanged factories)
# -----------------------------
def create_cnn_model(input_shape):
    model = Sequential([
        Conv1D(filters=64, kernel_size=3, activation='relu', input_shape=input_shape),
        BatchNormalization(),
        MaxPooling1D(pool_size=2),
        Dropout(0.3),
        Conv1D(filters=32, kernel_size=3, activation='relu'),
        BatchNormalization(),
        MaxPooling1D(pool_size=2),
        Dropout(0.3),
        Flatten(),
        Dense(64, activation='relu'),
        BatchNormalization(),
        Dropout(0.4),
        Dense(32, activation='relu'),
        Dropout(0.3),
        Dense(1, activation='linear')
    ])
    model.compile(optimizer=Adam(learning_rate=0.001), loss='mse', metrics=['mae'])
    return model

def create_hybrid_model(input_shape):
    cnn_input = Input(shape=input_shape, name='cnn_input')
    cnn_branch = Conv1D(filters=32, kernel_size=3, activation='relu')(cnn_input)
    cnn_branch = BatchNormalization()(cnn_branch)
    cnn_branch = MaxPooling1D(pool_size=2)(cnn_branch)
    cnn_branch = Dropout(0.3)(cnn_branch)
    cnn_branch = Flatten()(cnn_branch)
    lstm_branch = LSTM(32, return_sequences=True)(cnn_input)
    lstm_branch = Dropout(0.3)(lstm_branch)
    lstm_branch = LSTM(16)(lstm_branch)
    lstm_branch = Dropout(0.3)(lstm_branch)
    combined = Concatenate()([cnn_branch, lstm_branch])
    combined = Dense(32, activation='relu')(combined)
    combined = BatchNormalization()(combined)
    combined = Dropout(0.4)(combined)
    combined = Dense(16, activation='relu')(combined)
    combined = Dropout(0.3)(combined)
    output = Dense(1, activation='linear')(combined)
    model = Model(inputs=cnn_input, outputs=output)
    model.compile(optimizer=RMSprop(learning_rate=0.001), loss='mse', metrics=['mae'])
    return model

# -----------------------------
# Visualization functions (modified)
# -----------------------------
def create_model_performance_bar_chart(metrics_dict, output_path="model_performance.png"):
    models = list(metrics_dict.keys())
    r2_scores = [metrics_dict[model].get('R²', 0) for model in models]
    mae_scores = [metrics_dict[model].get('MAE', 0) for model in models]
    rmse_scores = [metrics_dict[model].get('RMSE', 0) for model in models]

    fig, (ax1, ax2, ax3) = plt.subplots(1, 3, figsize=(15, 5))

    bars1 = ax1.bar(models, r2_scores, alpha=0.9)
    ax1.set_title('R² Score Comparison')
    ax1.set_ylabel('R² Score')
    ax1.set_ylim(min(0, min(r2_scores) - 0.1), 1)
    ax1.tick_params(axis='x', rotation=45)
    for bar, score in zip(bars1, r2_scores):
        height = bar.get_height()
        ax1.text(bar.get_x() + bar.get_width()/2., height + 0.01, f'{score:.3f}', ha='center', va='bottom')

    bars2 = ax2.bar(models, mae_scores, alpha=0.9)
    ax2.set_title('Mean Absolute Error Comparison')
    ax2.set_ylabel('MAE')
    ax2.tick_params(axis='x', rotation=45)
    for bar, score in zip(bars2, mae_scores):
        height = bar.get_height()
        ax2.text(bar.get_x() + bar.get_width()/2., height + 0.001, f'{score:.3f}', ha='center', va='bottom')

    bars3 = ax3.bar(models, rmse_scores, alpha=0.9)
    ax3.set_title('RMSE Comparison')
    ax3.set_ylabel('RMSE')
    ax3.tick_params(axis='x', rotation=45)
    for bar, score in zip(bars3, rmse_scores):
        height = bar.get_height()
        ax3.text(bar.get_x() + bar.get_width()/2., height + 0.001, f'{score:.3f}', ha='center', va='bottom')

    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    return output_path

def create_prediction_bar_chart(predictions, future_years, output_path="predictions_comparison.png"):
    # include all models (including ARIMA if present)
    models = list(predictions.keys())
    years = future_years.flatten()
    fig, ax = plt.subplots(figsize=(14, 7))

    # Prepare grouped bars: each model gets its own grouped column per year
    x = np.arange(len(years))
    n_models = len(models)
    width = 0.8 / max(1, n_models)

    colors = plt.cm.tab20(np.linspace(0, 1, max(1, n_models)))
    for i, model_name in enumerate(models):
        pred_values = np.array(predictions[model_name], dtype=float)
        ax.bar(x + i * width, pred_values, width, label=model_name, alpha=0.9, color=colors[i % len(colors)])

    ax.set_xlabel('Future Years')
    ax.set_ylabel('Predicted GCI Values')
    ax.set_title('Model Predictions Comparison (All Models)')
    ax.set_xticks(x + width * (n_models - 1) / 2)
    ax.set_xticklabels([str(int(year)) for year in years])
    ax.legend(loc='best', fontsize=9)
    ax.grid(True, alpha=0.25)

    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    return output_path

def create_model_validation_chart(validation_results, output_path="model_validation.png"):
    models = list(validation_results.keys())
    cv_scores = [validation_results[m].get('cv_mean', 0) for m in models]
    train_scores = [validation_results[m].get('train_score', 0) for m in models]
    statuses = [validation_results[m].get('status', 'Unknown') for m in models]

    fig, ax = plt.subplots(figsize=(12, 6))
    x = np.arange(len(models))
    width = 0.35
    bars1 = ax.bar(x - width/2, cv_scores, width, label='CV Score', alpha=0.9)
    bars2 = ax.bar(x + width/2, train_scores, width, label='Train Score', alpha=0.9)

    for i, status in enumerate(statuses):
        color = 'green' if status == 'Good' else 'orange' if status == 'Needs Tuning' else 'blue' if status == 'Limited Data' else 'red'
        bars1[i].set_color(color)
        bars2[i].set_color(color)

    ax.set_xlabel('Models')
    ax.set_ylabel('R² Score')
    ax.set_title('Model Validation Results')
    ax.set_xticks(x)
    ax.set_xticklabels(models, rotation=45)
    ax.legend()
    ax.grid(True, alpha=0.3)
    ax.set_ylim(0, 1)
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    return output_path

def create_combined_time_series(years, historical_gci, future_years, predictions, uncertainties=None, output_path="combined_time_series.png"):
    plt.figure(figsize=(15, 9))

    # historical
    yrs_hist = [y for y, g in zip(years, historical_gci) if g is not None]
    gci_hist = [g for g in historical_gci if g is not None]
    if len(yrs_hist) >= 1:
        plt.plot(yrs_hist, gci_hist, 'o-', label='Historical GCI', linewidth=3, markersize=8, color='black')
        if len(yrs_hist) >= 2:
            z = np.polyfit(yrs_hist, gci_hist, 1)
            p = np.poly1d(z)
            plt.plot(yrs_hist, p(yrs_hist), '--', linewidth=1.5, label=f'Historical Trend (slope={z[0]:.4f})', color='gray')

    # plot each model prediction as simple line, no CI shading
    colors = plt.cm.tab10(np.linspace(0, 1, max(1, len(predictions))))
    model_names = list(predictions.keys())
    yrs_future = future_years.flatten()

    for i, name in enumerate(model_names):
        preds = np.array(predictions[name], dtype=float)
        plt.plot(yrs_future, preds, linestyle='--', marker='x', label=f'{name} Prediction', color=colors[i % len(colors)], linewidth=2)

    # risk zones
    plt.axhspan(0.0, 0.2, alpha=0.12, color='red', label='High Risk Zone (GCI < 0.2)')
    plt.axhspan(0.2, 0.4, alpha=0.10, color='orange', label='Medium Risk Zone (0.2 ≤ GCI < 0.4)')
    plt.axhspan(0.4, 1.0, alpha=0.08, color='green', label='Low Risk Zone (GCI ≥ 0.4)')

    max_hist_year = max(years) if len(years) > 0 else (min(yrs_future) - 1)
    sep_x = max_hist_year + 0.5
    plt.axvline(x=sep_x, color='gray', linestyle=':', linewidth=2)
    ymin, ymax = plt.ylim()
    plt.text(sep_x + 0.1, ymax - (ymax - ymin) * 0.05, 'Historical → Future', rotation=90, color='gray')

    plt.title("Combined GCI Time Series: Historical & Model Predictions", fontsize=16, fontweight='bold')
    plt.xlabel("Year", fontsize=14)
    plt.ylabel("GCI (NDVI)", fontsize=14)
    plt.grid(True, alpha=0.25)
    plt.legend(loc='best', fontsize=9, ncol=2)
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    return output_path

# -----------------------------
# Advanced ML models (unchanged except ensemble and n_jobs)
# -----------------------------
MIN_DL_SAMPLES = 100

def enhanced_models(X, y, environmental_features=None):
    models_dict = {}

    if environmental_features:
        enhanced_X = []
        for i, year_features in enumerate(X):
            env_feat = environmental_features[i] if i < len(environmental_features) else {}
            feature_vector = list(year_features) + [
                env_feat.get('temperature', 0) if env_feat.get('temperature') is not None else 0,
                env_feat.get('precipitation', 0) if env_feat.get('precipitation') is not None else 0,
                env_feat.get('urban_fraction', 0) if env_feat.get('urban_fraction') is not None else 0,
                env_feat.get('no2_concentration', 0) if env_feat.get('no2_concentration') is not None else 0
            ]
            enhanced_X.append(feature_vector)
        enhanced_X = np.array(enhanced_X)
    else:
        enhanced_X = np.array(X)

    enhanced_X = np.nan_to_num(enhanced_X, nan=0.0, posinf=0.0, neginf=0.0)

    scaler = StandardScaler()
    enhanced_X_scaled = scaler.fit_transform(enhanced_X)

    y_arr = np.array([v if v is not None else np.nan for v in y], dtype=float)
    mask = ~np.isnan(y_arr)
    if mask.sum() < 2:
        raise ValueError("Not enough valid GCI values to train models.")
    X_train = enhanced_X_scaled[mask]
    y_train = y_arr[mask]

    lr = LinearRegression()
    lr.fit(X_train, y_train)
    models_dict['Linear Regression'] = lr

    svr = SVR(kernel='rbf', C=100, gamma=0.1, epsilon=.01)
    svr.fit(X_train, y_train)
    models_dict['Support Vector Machine'] = svr

    rf = RandomForestRegressor(n_estimators=200, random_state=42, max_depth=10, n_jobs=-1)
    rf.fit(X_train, y_train)
    models_dict['Random Forest'] = rf

    gb = GradientBoostingRegressor(n_estimators=200, learning_rate=0.1, max_depth=6)
    gb.fit(X_train, y_train)
    models_dict['Gradient Boosting'] = gb

    try:
        ensemble = VotingRegressor([('lr', lr), ('rf', rf), ('gb', gb)])
        ensemble.fit(X_train, y_train)
        models_dict['Ensemble'] = ensemble
    except Exception as e:
        print(f"[WARN] Ensemble creation failed: {e}")

    if DL_AVAILABLE and len(X_train) >= MIN_DL_SAMPLES:
        try:
            X_reshaped = enhanced_X_scaled.reshape(enhanced_X_scaled.shape[0], enhanced_X_scaled.shape[1], 1)
            cnn_model = create_cnn_model((enhanced_X_scaled.shape[1], 1))
            cnn_model.fit(X_reshaped[mask], y_train, epochs=50, verbose=0, validation_split=0.2,
                         callbacks=[EarlyStopping(patience=5, restore_best_weights=True)])
            models_dict['CNN Deep Learning'] = cnn_model

            hybrid_model = create_hybrid_model((enhanced_X_scaled.shape[1], 1))
            hybrid_model.fit(X_reshaped[mask], y_train, epochs=50, verbose=0, validation_split=0.2,
                             callbacks=[EarlyStopping(patience=5, restore_best_weights=True)])
            models_dict['Hybrid CNN-LSTM'] = hybrid_model
        except Exception as e:
            print(f"[WARN] Deep Learning models failed: {e}")
    else:
        if DL_AVAILABLE:
            print(f"[WARN] DL available but insufficient samples ({len(X_train)}). Skipping DL models.")
        else:
            print("[INFO] DL not available; skipping DL models.")

    if ARIMA_AVAILABLE and len(y_train) >= 4:
        try:
            arima = ARIMA(y_train, order=(2,1,1))
            arima_fit = arima.fit()
            models_dict['ARIMA'] = arima_fit
        except Exception as e:
            print(f"[WARN] ARIMA model failed: {e}")

    return models_dict, scaler

def make_predictions(models_dict, scaler, future_years, environmental_features=None):
    predictions = {}
    future_X = create_enhanced_features(future_years.flatten(), [])

    if environmental_features:
        last_env_features = environmental_features[-1] if environmental_features else {}
    else:
        last_env_features = {'temperature': 0, 'precipitation': 0, 'urban_fraction': 0, 'no2_concentration': 0}

    enhanced_future_X = []
    for year_features in future_X:
        feature_vector = list(year_features) + [
            last_env_features.get('temperature', 0) if last_env_features.get('temperature') is not None else 0,
            last_env_features.get('precipitation', 0) if last_env_features.get('precipitation') is not None else 0,
            last_env_features.get('urban_fraction', 0) if last_env_features.get('urban_fraction') is not None else 0,
            last_env_features.get('no2_concentration', 0) if last_env_features.get('no2_concentration') is not None else 0
        ]
        enhanced_future_X.append(feature_vector)
    enhanced_future_X = np.array(enhanced_future_X)
    enhanced_future_X = np.nan_to_num(enhanced_future_X, nan=0.0, posinf=0.0, neginf=0.0)

    try:
        enhanced_future_X_scaled = scaler.transform(enhanced_future_X)
    except Exception:
        enhanced_future_X_scaled = enhanced_future_X

    for name, model in models_dict.items():
        try:
            if name == 'ARIMA':
                pred = model.forecast(steps=len(future_years))
                predictions[name] = np.array(pred).astype(float).flatten()
            elif name in ['CNN Deep Learning', 'Hybrid CNN-LSTM']:
                X_reshaped = enhanced_future_X_scaled.reshape(enhanced_future_X_scaled.shape[0], enhanced_future_X_scaled.shape[1], 1)
                pred = model.predict(X_reshaped).flatten()
                predictions[name] = np.array(pred).astype(float)
            else:
                pred = model.predict(enhanced_future_X_scaled)
                predictions[name] = np.array(pred).astype(float)
        except Exception as e:
            print(f"[WARN] Prediction failed for {name}: {e}")
            predictions[name] = np.zeros(len(future_years), dtype=float)
    return predictions

# -----------------------------
# Uncertainty quantification (unchanged but used only for numeric CI storage)
# -----------------------------
def get_prediction_intervals(models_dict, scaler, X, y, future_X, confidence=0.95, n_bootstrap=50):
    intervals = {}
    X_arr = np.array(X)
    y_arr = np.array(y, dtype=float)
    try:
        X_scaled = scaler.transform(X_arr)
    except Exception:
        X_scaled = X_arr
    try:
        future_scaled = scaler.transform(np.array(future_X))
    except Exception:
        future_scaled = np.array(future_X)

    if len(y_arr) < 2:
        return intervals

    for name, model in models_dict.items():
        if name == 'ARIMA':
            continue
        bootstrap_preds = []
        for _ in range(n_bootstrap):
            try:
                indices = np.random.choice(len(X_scaled), len(X_scaled), replace=True)
                X_boot = X_scaled[indices]
                y_boot = y_arr[indices]
                if name == 'Linear Regression':
                    m = LinearRegression()
                elif name == 'Support Vector Machine':
                    m = SVR(kernel='rbf', C=100, gamma=0.1, epsilon=.01)
                elif name == 'Random Forest':
                    m = RandomForestRegressor(n_estimators=100, random_state=42)
                elif name == 'Gradient Boosting':
                    m = GradientBoostingRegressor(n_estimators=100, learning_rate=0.1)
                elif name == 'Ensemble':
                    m = VotingRegressor([('lr', LinearRegression()), ('rf', RandomForestRegressor(n_estimators=50, random_state=42)), ('gb', GradientBoostingRegressor(n_estimators=50))])
                else:
                    continue
                m.fit(X_boot, y_boot)
                p = m.predict(future_scaled)
                bootstrap_preds.append(p)
            except Exception:
                continue
        if bootstrap_preds:
            bp = np.array(bootstrap_preds)
            lower = np.percentile(bp, (1-confidence)/2 * 100, axis=0)
            upper = np.percentile(bp, (1+confidence)/2 * 100, axis=0)
            intervals[name] = {'lower': lower, 'upper': upper}
    return intervals

# -----------------------------
# Metrics calculation
# -----------------------------
def calculate_enhanced_model_metrics(years, gci_values, models_dict, scaler, environmental_features=None):
    X = create_enhanced_features(years, [])
    if environmental_features:
        enhanced_X = []
        for i, year_features in enumerate(X):
            env_feat = environmental_features[i] if i < len(environmental_features) else {}
            feature_vector = list(year_features) + [
                env_feat.get('temperature', 0) if env_feat.get('temperature') is not None else 0,
                env_feat.get('precipitation', 0) if env_feat.get('precipitation') is not None else 0,
                env_feat.get('urban_fraction', 0) if env_feat.get('urban_fraction') is not None else 0,
                env_feat.get('no2_concentration', 0) if env_feat.get('no2_concentration') is not None else 0
            ]
            enhanced_X.append(feature_vector)
        enhanced_X = np.array(enhanced_X)
    else:
        enhanced_X = X

    enhanced_X = np.nan_to_num(enhanced_X, nan=0.0)
    try:
        enhanced_X_scaled = scaler.transform(enhanced_X)
    except Exception:
        enhanced_X_scaled = enhanced_X

    y = np.array([v if v is not None else np.nan for v in gci_values])
    mask = ~np.isnan(y)
    X_valid = enhanced_X_scaled[mask]
    y_valid = y[mask]

    metrics_dict = {}
    for name, model in models_dict.items():
        if name == 'ARIMA':
            continue
        try:
            y_pred = model.predict(X_valid)
            mse = mean_squared_error(y_valid, y_pred)
            rmse = np.sqrt(mse)
            metrics_dict[name] = {
                'R²': r2_score(y_valid, y_pred),
                'MAE': mean_absolute_error(y_valid, y_pred),
                'RMSE': rmse,
                'MAPE': np.mean(np.abs((y_valid - y_pred) / (y_valid + 1e-8))) * 100
            }
        except Exception as e:
            print(f"[WARN] Metrics calculation failed for {name}: {e}")
            metrics_dict[name] = {'R²': 0, 'MAE': 0, 'RMSE': 0, 'MAPE': 0}
    return metrics_dict

# -----------------------------
# Report generation (improved: include metrics table after performance chart and all model predictions plotted)
# -----------------------------
def save_enhanced_report_with_risks(file_path, years, gci_values, future_years, models_dict,
                                  metrics_dict, plot_points, area, within_buffer, distance,
                                  closest_lake, risks, environmental_data, uncertainties, scaler,
                                  validation_results=None, output_file="report.docx"):
    doc = Document()
    doc.add_heading("Enhanced Green Cover Index (GCI) Comprehensive Report", level=1)
    doc.add_paragraph("Generated by: Automated GCI Analyzer")
    doc.add_paragraph("This report contains historical GCI (NDVI) trends, model performance, and future predictions for the analysed plot.")
    doc.add_page_break()

    # Plot Analysis
    doc.add_heading("1. Plot Analysis", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Latitude"
    hdr[1].text = "Longitude"
    for lat, lon in plot_points:
        row = table.add_row().cells
        row[0].text = f"{lat:.6f}"
        row[1].text = f"{lon:.6f}"
    doc.add_paragraph(f"Calculated area: {area:.2f} m² ({area/10000:.2f} hectares)")
    if within_buffer:
        doc.add_paragraph(f"Plot is within legal buffer, {distance:.2f} m from nearest lake.")
    else:
        doc.add_paragraph("Plot outside legal buffer. Closest lake details:")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Attribute"
        hdr[1].text = "Value"
        for k, v in (closest_lake or {}).items():
            row = table.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)

    doc.add_page_break()

    # Environmental Risk Assessment
    doc.add_heading("2. Environmental Risk Assessment", level=2)
    doc.add_paragraph("Risk assessment summary based on historical GCI and environmental indicators.")
    risk_table = doc.add_table(rows=1, cols=2)
    risk_table.style = "Table Grid"
    hdr = risk_table.rows[0].cells
    hdr[0].text = "Risk Category"
    hdr[1].text = "Risk Level"
    for risk_type, risk_level in risks.items():
        row = risk_table.add_row().cells
        row[0].text = risk_type.replace('_', ' ').title()
        row[1].text = risk_level

    doc.add_page_break()

    # Historical GCI
    doc.add_heading("3. Historical GCI Data & Trend", level=2)
    hist_table = doc.add_table(rows=1, cols=2)
    hist_table.style = "Table Grid"
    hdr = hist_table.rows[0].cells
    hdr[0].text = "Year"
    hdr[1].text = "GCI"
    for y, gci in zip(years, gci_values):
        row = hist_table.add_row().cells
        row[0].text = str(y)
        row[1].text = "N/A" if gci is None else f"{gci:.4f}"

    # Historical plot image
    hist_img = "enhanced_historical_gci.png"
    try:
        yrs_plot = [y for y, g in zip(years, gci_values) if g is not None]
        gci_plot = [g for g in gci_values if g is not None]
        plt.figure(figsize=(10, 6))
        if len(yrs_plot) > 0:
            plt.plot(yrs_plot, gci_plot, marker='o', linewidth=2, markersize=6, color='black', label='Historical GCI')
            if len(yrs_plot) > 1:
                z = np.polyfit(yrs_plot, gci_plot, 1)
                p = np.poly1d(z)
                plt.plot(yrs_plot, p(yrs_plot), '--', alpha=0.8, label=f'Trend (slope: {z[0]:.4f})', color='gray')
        plt.title("Historical GCI (NDVI)")
        plt.xlabel("Year"); plt.ylabel("GCI (NDVI)")
        plt.grid(alpha=0.25); plt.legend()
        plt.tight_layout(); plt.savefig(hist_img, dpi=300); plt.close()
        doc.add_picture(hist_img, width=Inches(6)); os.remove(hist_img)
    except Exception as e:
        print(f"[WARN] Could not create/attach historical plot: {e}")

    doc.add_page_break()

    # Model performance bar chart + metrics table right after
    doc.add_heading("4. Model Performance", level=2)
    perf_img = "model_performance.png"
    try:
        if metrics_dict:
            create_model_performance_bar_chart(metrics_dict, perf_img)
            doc.add_paragraph("Model performance visualisation (R², MAE, RMSE).")
            doc.add_picture(perf_img, width=Inches(6))
            os.remove(perf_img)

            # Insert metrics table directly after the bar chart
            doc.add_paragraph("Detailed metrics table:")
            metrics_table = doc.add_table(rows=1, cols=5)
            metrics_table.style = "Table Grid"
            hdr = metrics_table.rows[0].cells
            hdr[0].text = "Model"
            hdr[1].text = "R²"
            hdr[2].text = "MAE"
            hdr[3].text = "RMSE"
            hdr[4].text = "MAPE (%)"
            # sort by R² descending for readability
            sorted_models = sorted(metrics_dict.items(), key=lambda x: x[1].get('R²', 0), reverse=True)
            for name, m in sorted_models:
                row = metrics_table.add_row().cells
                row[0].text = name
                row[1].text = f"{m.get('R²', 0):.3f}"
                row[2].text = f"{m.get('MAE', 0):.3f}"
                row[3].text = f"{m.get('RMSE', 0):.3f}"
                row[4].text = f"{m.get('MAPE', 0):.1f}"
    except Exception as e:
        print(f"[WARN] Could not create/attach performance image or metrics table: {e}")

    doc.add_page_break()

    # Predictions (bar chart of all models)
    doc.add_heading("5. Future Predictions (All Models)", level=2)
    pred_bar_img = "predictions_comparison.png"
    try:
        predictions = make_predictions(models_dict, scaler, future_years, environmental_data)
        create_prediction_bar_chart(predictions, future_years, pred_bar_img)
        doc.add_paragraph("Predictions comparison across all trained models (bar chart).")
        doc.add_picture(pred_bar_img, width=Inches(6))
        os.remove(pred_bar_img)
    except Exception as e:
        print(f"[WARN] Could not create/attach predictions bar chart: {e}")

    doc.add_page_break()

    # Combined time series (professional, all models, no CI shading)
    doc.add_heading("6. Combined Time Series: Historical + All Model Predictions", level=2)
    combined_img = "combined_time_series.png"
    try:
        predictions = make_predictions(models_dict, scaler, future_years, environmental_data)
        create_combined_time_series(years, gci_values, future_years, predictions, uncertainties, combined_img)
        doc.add_paragraph("Comprehensive time series that overlays historical GCI and each model's predictions.")
        doc.add_picture(combined_img, width=Inches(8))
        os.remove(combined_img)
    except Exception as e:
        print(f"[WARN] Could not create/attach combined time series: {e}")

    doc.add_page_break()

    # Predictions tables (per model) - CI column kept but filled N/A since user requested no CI
    doc.add_heading("7. Predictions Table (per model)", level=2)
    try:
        for model_name, pred_values in predictions.items():
            doc.add_heading(f"{model_name} Predictions", level=3)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Year"
            hdr[1].text = "Predicted GCI"
            for i, (year, pred) in enumerate(zip(future_years.flatten(), pred_values)):
                row = table.add_row().cells
                row[0].text = str(int(year))
                row[1].text = f"{pred:.4f}"
                
    except Exception as e:
        print(f"[WARN] Could not populate predictions tables: {e}")

    # Validation table (if present)
    if validation_results:
        doc.add_page_break()
        doc.add_heading("8. Model Validation and Tuning Assessment", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Model"; hdr[1].text = "CV Score"; hdr[2].text = "Train Score"; hdr[3].text = "Status"; hdr[4].text = "Recommendations"
        for model_name, results in validation_results.items():
            row = table.add_row().cells
            row[0].text = model_name
            row[1].text = f"{results.get('cv_mean', 0):.3f} ± {results.get('cv_std', 0):.3f}"
            row[2].text = f"{results.get('train_score', 0):.3f}"
            row[3].text = results.get('status', 'Unknown')
            if results.get('status') == 'Good':
                row[4].text = "Model is well-tuned and performing optimally"
            elif results.get('status') == 'Needs Tuning':
                row[4].text = "Consider hyperparameter optimisation & feature engineering"
            elif results.get('status') == 'Limited Data':
                row[4].text = "Limited data - collect more historical points"
            else:
                row[4].text = "Manual review recommended"

    # Recommendations
    doc.add_page_break()
    doc.add_heading("9. Recommendations", level=2)
    if risks.get('vegetation_decline') == 'High':
        doc.add_paragraph("• HIGH PRIORITY: Implement immediate vegetation restoration measures")
    if risks.get('climate_vulnerability') == 'High':
        doc.add_paragraph("• HIGH PRIORITY: Develop climate adaptation strategies")
    if risks.get('water_stress') == 'High':
        doc.add_paragraph("• HIGH PRIORITY: Implement water conservation measures")
    if risks.get('urbanization_pressure') == 'High':
        doc.add_paragraph("• MEDIUM PRIORITY: Monitor urban development impacts")
    if risks.get('air_quality') == 'High':
        doc.add_paragraph("• MEDIUM PRIORITY: Address air quality concerns")
    doc.add_paragraph("• Regular monitoring and model updates recommended")
    doc.add_paragraph("• Consider adaptive management & targeted field surveys to validate predictions")

    doc.save(output_file)
    print(f"Enhanced report saved as {output_file}")

# -----------------------------
# Main pipeline
# -----------------------------
def enhanced_main(file_path, lake_csv):
    print("Initializing Earth Engine...")
    ee_ok = initialize_earth_engine()
    if not ee_ok:
        print("[ERROR] Earth Engine not available. The script will continue but environmental features may be None.")

    print("Processing plot data...")
    plot_points, plot_points_gci = get_plot_points_from_csv(file_path)
    if not plot_points:
        return print("No valid plot points.")

    area = calculate_polygon_area(plot_points)
    print(f"Plot area: {area:.2f} m² ({area/10000:.2f} hectares)")

    print("Processing lake data...")
    lakes = parse_csv_lakes(lake_csv)
    lake_coords = [(lake['latitude'], lake['longitude']) for lake in lakes] if lakes else []
    within_buffer, distance, closest_lake = is_within_buffer(plot_points, lake_coords, lakes) if lakes else (False, float('inf'), None)
    print(f"Buffer analysis: Within buffer = {within_buffer}, Distance = {distance:.2f}m")

    # GCI and environmental features
    lat = np.mean([p[0] for p in plot_points])
    lon = np.mean([p[1] for p in plot_points])
    city = "Bangalore"

    years = np.arange(2015, 2025)
    print("Fetching GCI (NDVI) for each historical year (cached)...")
    gci_values = calculate_gci_values_from_gee(lat, lon, years)
    print("Fetching environmental features for each year (cached)...")
    environmental_data = [get_environmental_features(lat, lon, city, int(year)) for year in years]

    print(f"GCI values (sample): {gci_values[:5]}")
    print(f"Environmental sample (latest): {environmental_data[-1] if environmental_data else None}")

    print("Performing risk assessment...")
    risks = calculate_risk_metrics(plot_points_gci, gci_values, environmental_data)
    print(f"Risk assessment: {risks}")

    print("Training models (speed tuned)...")
    X = create_enhanced_features(years, [])
    try:
        models_dict, scaler = enhanced_models(X, [v if v is not None else np.nan for v in gci_values], environmental_data)
    except Exception as e:
        print("[ERROR] Model training failed:", e)
        return

    print("Calculating metrics...")
    metrics_dict = calculate_enhanced_model_metrics(years, gci_values, models_dict, scaler, environmental_data)

    print("Validating model performance...")
    validation_results = {}
    enhanced_X = []
    for i, year_features in enumerate(X):
        env_feat = environmental_data[i] if i < len(environmental_data) else {}
        feature_vector = list(year_features) + [
            env_feat.get('temperature', 0) if env_feat.get('temperature') is not None else 0,
            env_feat.get('precipitation', 0) if env_feat.get('precipitation') is not None else 0,
            env_feat.get('urban_fraction', 0) if env_feat.get('urban_fraction') is not None else 0,
            env_feat.get('no2_concentration', 0) if env_feat.get('no2_concentration') is not None else 0
        ]
        enhanced_X.append(feature_vector)
    enhanced_X = np.array(enhanced_X)
    try:
        enhanced_X_scaled = scaler.transform(enhanced_X)
    except Exception:
        enhanced_X_scaled = enhanced_X
    y_arr = np.array([v if v is not None else np.nan for v in gci_values])
    mask = ~np.isnan(y_arr)
    X_valid = enhanced_X_scaled[mask]
    y_valid = y_arr[mask]
    for model_name, model in models_dict.items():
        if model_name not in ['ARIMA', 'CNN Deep Learning', 'Hybrid CNN-LSTM']:
            try:
                validation_results[model_name] = validate_model_performance(model, X_valid, y_valid, model_name)
            except Exception as e:
                validation_results[model_name] = {'error': str(e), 'status': 'Error'}

    print("Making future predictions...")
    future_years = np.arange(2025, 2036).reshape(-1, 1)
    predictions = make_predictions(models_dict, scaler, future_years, environmental_data)

    print("Calculating uncertainties (bootstrap)...")
    try:
        uncertainties = get_prediction_intervals(models_dict, scaler, enhanced_X[mask], y_valid, create_enhanced_features(future_years.flatten(), []))
    except Exception:
        uncertainties = {}

    print("Generating report (may still take time for plotting/reporting)...")
    save_enhanced_report_with_risks(file_path, years, gci_values, future_years, models_dict,
                                  metrics_dict, plot_points, area, within_buffer, distance,
                                  closest_lake, risks, environmental_data, uncertainties, scaler,
                                  validation_results)
    print("[INFO] Done.")

# -----------------------------
# Entrypoint
# -----------------------------
if __name__ == "__main__":
    plot_csv = sys.argv[1] if len(sys.argv) > 1 else "coordinates.csv"
    lake_csv = sys.argv[2] if len(sys.argv) > 2 else "output_lake.csv"
    enhanced_main(plot_csv, lake_csv)
