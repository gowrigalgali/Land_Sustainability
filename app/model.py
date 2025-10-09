import math
import csv
import os
import numpy as np
import matplotlib.pyplot as plt
from geopy.distance import geodesic
import ee
from docx import Document
from docx.shared import Inches

# ML Models
from sklearn.linear_model import LinearRegression
from sklearn.svm import SVR
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error

# Optional Deep Learning
try:
    from tensorflow.keras.models import Sequential
    from tensorflow.keras.layers import Dense
    from sklearn.preprocessing import MinMaxScaler
    DL_AVAILABLE = True
except ImportError:
    DL_AVAILABLE = False

# -----------------------------
# Earth Engine Authentication
# -----------------------------
ee.Authenticate()
ee.Initialize(project='crisiscompass')

# -----------------------------
# Polygon Area Calculation
# -----------------------------
def calculate_polygon_area(vertices):
    R = 6378137
    vertices = [(math.radians(lat), math.radians(lon)) for lat, lon in vertices]
    area = 0
    for i in range(len(vertices)):
        lat1, lon1 = vertices[i]
        lat2, lon2 = vertices[(i + 1) % len(vertices)]
        area += (lon2 - lon1) * (2 + math.sin(lat1) + math.sin(lat2))
    return abs(area) / 2.0 * (R**2)

# -----------------------------
# CSV Parsing for Lakes
# -----------------------------
def parse_csv_lakes(file_path):
    lakes = []
    with open(file_path, "r", newline='', encoding='utf-8') as csv_file:
        reader = csv.DictReader(csv_file)
        for row in reader:
            lat, lon = row.get('Latitude'), row.get('Longitude')
            if lat and lon:
                try:
                    lake = {k: row[k] for k in row.keys()}
                    lake['latitude'] = float(lat)
                    lake['longitude'] = float(lon)
                    lakes.append(lake)
                except ValueError:
                    continue
    return lakes

# -----------------------------
# Buffer Check
# -----------------------------
def is_within_buffer(plot_vertices, lake_coords, lakes, min_buffer=30, max_buffer=75):
    plot_center = (sum([lat for lat, _ in plot_vertices])/len(plot_vertices),
                   sum([lon for _, lon in plot_vertices])/len(plot_vertices))
    closest_lake = None
    closest_distance = float('inf')
    for i, lake_coord in enumerate(lake_coords):
        distance = geodesic(plot_center, lake_coord).meters
        if distance < closest_distance:
            closest_distance = distance
            closest_lake = lakes[i]
        if min_buffer <= distance <= max_buffer:
            return True, distance, closest_lake
    return False, closest_distance, closest_lake

# -----------------------------
# Get Closest Lake Details
# -----------------------------
def get_closest_lake_details(closest_lake):
    if closest_lake:
        for key, value in closest_lake.items():
            print(f"{key}: {value}")
    else:
        print("No closest lake found.")

# -----------------------------
# NDVI / GCI Calculation
# -----------------------------
def calculate_ndvi(image):
    ndvi = image.normalizedDifference(['B8', 'B4']).rename('NDVI')
    return image.addBands(ndvi)

def calculate_current_gci(coordinates):
    if coordinates[0] != coordinates[-1]:
        coordinates.append(coordinates[0])
    aoi = ee.Geometry.Polygon([coordinates])
    image_collection = ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED') \
        .filterBounds(aoi) \
        .filterDate('2025-01-20', '2025-01-25') \
        .map(calculate_ndvi) \
        .select('NDVI')
    mean_ndvi = image_collection.mean().reduceRegion(
        reducer=ee.Reducer.mean(), geometry=aoi, scale=30, maxPixels=1e8
    )
    return mean_ndvi.get('NDVI').getInfo()

def calculate_yearly_gci(coordinates):
    if coordinates[0] != coordinates[-1]:
        coordinates.append(coordinates[0])
    aoi = ee.Geometry.Polygon([coordinates])
    years = np.arange(2017, 2025)
    gci_values = []
    for year in years:
        year_start = f'{year}-04-01'
        year_end = f'{year+1}-04-01'
        image_collection = ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED') \
            .filterBounds(aoi) \
            .filterDate(year_start, year_end) \
            .map(calculate_ndvi) \
            .select('NDVI')
        mean_ndvi = image_collection.mean().reduceRegion(
            reducer=ee.Reducer.mean(), geometry=aoi, scale=30, maxPixels=1e8
        )
        gci_values.append(mean_ndvi.get('NDVI').getInfo())
    return years, gci_values

# -----------------------------
# Read Plot Points from CSV
# -----------------------------
def get_plot_points_from_csv(file_path):
    plot_points, plot_points_gci = [], []
    with open(file_path, "r", newline='', encoding='utf-8') as csv_file:
        reader = csv.DictReader(csv_file)
        for row in reader:
            lat, lon = row.get('Latitude'), row.get('Longitude')
            if lat and lon:
                try:
                    lat, lon = float(lat), float(lon)
                    plot_points.append((lat, lon))
                    plot_points_gci.append((lon, lat))
                except ValueError:
                    continue
    return plot_points, plot_points_gci

# -----------------------------
# Machine Learning Models
# -----------------------------
def fit_and_predict_models(years, gci_values, future_years):
    models_dict = {}
    X = np.array(years).reshape(-1, 1)
    y = np.array(gci_values)

    # Linear Regression
    lr = LinearRegression(); lr.fit(X,y)
    models_dict['Linear Regression'] = lr.predict(future_years)

    # SVR
    svr = SVR(kernel='rbf', C=100, gamma=0.1, epsilon=.01); svr.fit(X,y)
    models_dict['Support Vector Machine'] = svr.predict(future_years)

    # Random Forest
    rf = RandomForestRegressor(n_estimators=200, random_state=42); rf.fit(X,y)
    models_dict['Random Forest'] = rf.predict(future_years)

    # Deep Learning
    if DL_AVAILABLE:
        scaler = MinMaxScaler()
        X_scaled = scaler.fit_transform(X)
        future_scaled = scaler.transform(future_years)
        model = Sequential([
            Dense(32, input_dim=1, activation='relu'),
            Dense(16, activation='relu'),
            Dense(1)
        ])
        model.compile(optimizer='adam', loss='mse')
        model.fit(X_scaled, y, epochs=500, verbose=0)
        models_dict['Deep Learning'] = model.predict(future_scaled).flatten()
    return models_dict

def calculate_model_metrics(years, gci_values, models_dict):
    X = np.array(years).reshape(-1,1)
    y = np.array(gci_values)
    metrics_dict = {}
    for name, predictions in models_dict.items():
        if name == 'Deep Learning': continue
        if name=='Linear Regression': model = LinearRegression().fit(X,y)
        elif name=='Support Vector Machine': model = SVR(kernel='rbf', C=100, gamma=0.1, epsilon=.01).fit(X,y)
        else: model = RandomForestRegressor(n_estimators=200, random_state=42).fit(X,y)
        y_pred = model.predict(X)
        mse = mean_squared_error(y, y_pred)
        rmse = np.sqrt(mse)
        metrics_dict[name] = {'R²': r2_score(y, y_pred),
                              'MAE': mean_absolute_error(y, y_pred),
                              'RMSE': rmse}
    return metrics_dict

# -----------------------------
# Generate Report
# -----------------------------
def save_enhanced_report(file_path, years, gci_values, future_years, models_dict, metrics_dict, plot_points, area, within_buffer, distance, closest_lake, output_file="report.docx"):
    doc = Document()
    doc.add_heading("Green Cover Index (GCI) Comprehensive Report", level=1)
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("This report provides analysis of GCI trends, model predictions, plot details, and lake proximity.")

    # Plot Analysis
    doc.add_heading("Plot Analysis", level=2)
    table = doc.add_table(rows=1, cols=2); table.style="Table Grid"
    hdr = table.rows[0].cells; hdr[0].text="Latitude"; hdr[1].text="Longitude"
    for lat, lon in plot_points:
        row = table.add_row().cells; row[0].text=f"{lat:.6f}"; row[1].text=f"{lon:.6f}"
    doc.add_paragraph(f"Calculated area: {area:.2f} m²")
    if within_buffer: doc.add_paragraph(f"Plot is within legal buffer, {distance:.2f} m from nearest lake.")
    else:
        doc.add_paragraph("Plot outside legal buffer. Closest lake details:")
        table = doc.add_table(rows=1, cols=2); table.style="Table Grid"
        hdr = table.rows[0].cells; hdr[0].text="Attribute"; hdr[1].text="Value"
        for k,v in closest_lake.items():
            row = table.add_row().cells; row[0].text=str(k); row[1].text=str(v)

    # Historical GCI
    doc.add_heading("Historical GCI Data", level=2)
    table = doc.add_table(rows=1, cols=2); table.style="Table Grid"
    hdr = table.rows[0].cells; hdr[0].text="Year"; hdr[1].text="GCI"
    for y,gci in zip(years,gci_values):
        row = table.add_row().cells; row[0].text=str(y); row[1].text=f"{gci:.2f}"
    plt.figure(figsize=(10,6)); plt.plot(years,gci_values,'o-',label='Historical GCI'); plt.title("Historical GCI"); plt.xlabel("Year"); plt.ylabel("GCI"); plt.grid(True); plt.legend(); plt.savefig("historical_gci.png"); plt.close()
    doc.add_picture("historical_gci.png", width=Inches(5)); os.remove("historical_gci.png")

    # Metrics Table
    doc.add_heading("Model Performance Metrics", level=2)
    table = doc.add_table(rows=1, cols=4); table.style="Table Grid"
    hdr=table.rows[0].cells; hdr[0].text="Model"; hdr[1].text="R²"; hdr[2].text="MAE"; hdr[3].text="RMSE"
    for name, m in metrics_dict.items():
        row = table.add_row().cells; row[0].text=name; row[1].text=f"{m['R²']:.3f}"; row[2].text=f"{m['MAE']:.3f}"; row[3].text=f"{m['RMSE']:.3f}"

    # Predictions
    doc.add_heading("Predictions (2025-2035)", level=2)
    for model_name, predictions in models_dict.items():
        doc.add_heading(model_name, level=3)
        table = doc.add_table(rows=1, cols=2); table.style="Table Grid"
        hdr=table.rows[0].cells; hdr[0].text="Year"; hdr[1].text="Predicted GCI"
        for y,gci in zip(future_years.flatten(), predictions):
            row=table.add_row().cells; row[0].text=str(int(y)); row[1].text=f"{gci:.3f}"
        # Plot each model
        plt.figure(figsize=(10,6)); plt.plot(years,gci_values,'o-',label='Historical GCI'); plt.plot(future_years,predictions,'x--',label=f'{model_name} Prediction')
        plt.title(f"GCI Predictions using {model_name}"); plt.xlabel("Year"); plt.ylabel("GCI"); plt.grid(True); plt.legend()
        fname=f"{model_name.replace(' ','_')}_pred.png"; plt.savefig(fname); plt.close(); doc.add_picture(fname,width=Inches(5)); os.remove(fname)

    # Comparative Plot
    plt.figure(figsize=(12,6)); plt.plot(years,gci_values,'o-',label='Historical GCI')
    for name, predictions in models_dict.items(): plt.plot(future_years,predictions,'--',marker='x',label=name)
    plt.title("Comparative GCI Predictions"); plt.xlabel("Year"); plt.ylabel("GCI"); plt.grid(True); plt.legend(); plt.savefig("comparative_gci.png"); plt.close(); doc.add_picture("comparative_gci.png",width=Inches(6)); os.remove("comparative_gci.png")

    doc.save(output_file)
    print(f"Report saved as {output_file}")

# -----------------------------
# Main Execution
# -----------------------------
def main(file_path, lake_csv):
    plot_points, plot_points_gci = get_plot_points_from_csv(file_path)
    if not plot_points: return print("No valid plot points.")
    area = calculate_polygon_area(plot_points)

    # Lake data
    lakes = parse_csv_lakes(lake_csv)
    lake_coords = [(lake['latitude'], lake['longitude']) for lake in lakes]
    within_buffer, distance, closest_lake = is_within_buffer(plot_points, lake_coords, lakes)

    # GCI Calculation
    current_gci = calculate_current_gci(plot_points_gci)
    print(f"Current GCI: {current_gci}")
    years, gci_values = calculate_yearly_gci(plot_points_gci)

    # Prediction Models
    future_years = np.arange(2025,2036).reshape(-1,1)
    models_dict = fit_and_predict_models(years, gci_values, future_years)
    metrics_dict = calculate_model_metrics(years, gci_values, models_dict)

    # Save Report
    save_enhanced_report(file_path, years, gci_values, future_years, models_dict, metrics_dict, plot_points, area, within_buffer, distance, closest_lake)

# -----------------------------
# Run
# -----------------------------
if __name__ == "__main__":
    plot_csv = "coordinates.csv"  # Your plot CSV
    lake_csv = "output_lake.csv"   # Lake details CSV
    main(plot_csv, lake_csv)
